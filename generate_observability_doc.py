"""
Team G Backend - Observability 구현 문서 생성 스크립트
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_document():
    doc = Document()

    # =========================================================================
    # 제목
    # =========================================================================
    title = doc.add_heading('Team G Backend - Observability 구현 문서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('로깅(Logging), 메트릭(Metrics), 트레이싱(Tracing) 구현 상세 설명')
    doc.add_paragraph('')

    # =========================================================================
    # 1. 개요
    # =========================================================================
    doc.add_heading('1. 개요', level=1)

    doc.add_paragraph(
        'Observability(관측 가능성)는 시스템의 내부 상태를 외부에서 관측할 수 있게 해주는 '
        '특성입니다. 현대 마이크로서비스 아키텍처에서는 세 가지 핵심 요소(Three Pillars)를 '
        '통해 시스템을 모니터링합니다:'
    )

    # 세 가지 기둥 설명
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # 헤더
    header_cells = table.rows[0].cells
    header_cells[0].text = '요소'
    header_cells[1].text = '설명'
    header_cells[2].text = '예시'

    # Metrics
    row = table.rows[1].cells
    row[0].text = 'Metrics (메트릭)'
    row[1].text = '수치로 표현되는 정량적 데이터. 시간에 따른 추세 파악에 적합'
    row[2].text = 'API 응답시간, 처리량, 에러율, CPU 사용률'

    # Logs
    row = table.rows[2].cells
    row[0].text = 'Logs (로그)'
    row[1].text = '이벤트 발생 시점의 상세한 텍스트 기록'
    row[2].text = '에러 메시지, 디버그 정보, 사용자 액션 기록'

    # Traces
    row = table.rows[3].cells
    row[0].text = 'Traces (트레이스)'
    row[1].text = '분산 시스템에서 요청이 여러 서비스를 거치는 전체 흐름 추적'
    row[2].text = 'API → Celery → 외부 API 호출 순서와 소요시간'

    doc.add_paragraph('')

    # =========================================================================
    # 2. 기술 스택 개요
    # =========================================================================
    doc.add_heading('2. 기술 스택 개요', level=1)

    doc.add_heading('2.1 전체 아키텍처', level=2)

    doc.add_paragraph(
        '본 프로젝트에서는 Grafana를 중심으로 세 가지 관측 데이터를 통합하여 시각화합니다.'
    )

    # 기술 스택 표
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '구분'
    header[1].text = '데이터 수집'
    header[2].text = '데이터 저장'
    header[3].text = '시각화'

    row = table.rows[1].cells
    row[0].text = 'Metrics'
    row[1].text = 'prometheus_client\ndjango_prometheus'
    row[2].text = 'Prometheus'
    row[3].text = 'Grafana'

    row = table.rows[2].cells
    row[0].text = 'Logs'
    row[1].text = 'Python logging\nlogging_loki'
    row[2].text = 'Loki'
    row[3].text = 'Grafana'

    row = table.rows[3].cells
    row[0].text = 'Traces'
    row[1].text = 'OpenTelemetry SDK\n(자동 계측)'
    row[2].text = 'Jaeger'
    row[3].text = 'Grafana / Jaeger UI'

    doc.add_paragraph('')

    doc.add_heading('2.2 사용 포트', level=2)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '서비스'
    header[1].text = '포트'
    header[2].text = '용도'

    ports = [
        ('Prometheus', '9090', '메트릭 저장소'),
        ('Pushgateway', '9091', 'Celery 메트릭 push'),
        ('Grafana', '3000', '통합 대시보드'),
        ('Loki', '3100', '로그 저장소'),
        ('Jaeger UI', '16686', '트레이스 조회'),
        ('Jaeger Agent', '6831', '트레이스 수집 (UDP)'),
        ('Django', '8000', '웹 API + /metrics'),
    ]

    for i, (service, port, desc) in enumerate(ports, 1):
        row = table.rows[i].cells
        row[0].text = service
        row[1].text = port
        row[2].text = desc

    doc.add_paragraph('')

    # =========================================================================
    # 3. 로깅 (Logging) 구현
    # =========================================================================
    doc.add_heading('3. 로깅 (Logging) 구현', level=1)

    doc.add_heading('3.1 기술 스택', level=2)

    doc.add_paragraph('• Python 내장 logging 모듈: 로그 생성의 기본 프레임워크')
    doc.add_paragraph('• logging_loki (python-logging-loki): Loki로 로그 전송하는 핸들러')
    doc.add_paragraph('• Loki: 로그 저장소 (Grafana Labs 제공)')
    doc.add_paragraph('• Grafana: 로그 시각화 및 검색')

    doc.add_heading('3.2 동작 원리', level=2)

    doc.add_paragraph(
        'Python logging 모듈은 "핸들러(Handler)" 패턴을 사용합니다. '
        '로그가 발생하면 등록된 모든 핸들러에게 로그가 전달되며, '
        '각 핸들러는 자신만의 방식으로 로그를 처리합니다.'
    )

    doc.add_paragraph('')
    doc.add_paragraph('로그 흐름:')
    doc.add_paragraph('logger.info("메시지") → [ConsoleHandler] → 터미널 출력')
    doc.add_paragraph('                     → [FileHandler] → 파일 저장')
    doc.add_paragraph('                     → [LokiHandler] → HTTP로 Loki 전송')
    doc.add_paragraph('')

    doc.add_heading('3.3 구현 코드 설명 (config/settings.py)', level=2)

    doc.add_heading('3.3.1 Loki 환경변수 설정', level=3)

    doc.add_paragraph('# config/settings.py (라인 258-260)')
    code_block = doc.add_paragraph()
    code_block.add_run('''LOKI_URL = os.getenv('LOKI_URL', 'http://localhost:3100/loki/api/v1/push')
LOKI_ENABLED = os.getenv('LOKI_ENABLED', 'true').lower() == 'true'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• LOKI_URL: Loki 서버의 push API 엔드포인트')
    doc.add_paragraph('• LOKI_ENABLED: 환경변수로 Loki 핸들러 활성화/비활성화 제어')
    doc.add_paragraph('')

    doc.add_heading('3.3.2 핸들러 정의', level=3)

    doc.add_paragraph('# config/settings.py (라인 262-291)')
    code_block = doc.add_paragraph()
    code_block.add_run('''_log_handlers = {
    'console': {
        'class': 'logging.StreamHandler',    # 터미널 출력
        'formatter': 'verbose',
    },
    'file': {
        'class': 'logging.handlers.RotatingFileHandler',  # 파일 저장
        'filename': LOG_DIR / 'django.log',
        'maxBytes': 10 * 1024 * 1024,  # 10MB 초과시 로테이션
        'backupCount': 5,               # 백업 파일 5개 유지
        'formatter': 'json',
    },
}

# Loki 핸들러 조건부 추가
if LOKI_ENABLED:
    try:
        import logging_loki
        _log_handlers['loki'] = {
            'class': 'logging_loki.LokiHandler',
            'url': LOKI_URL,
            'tags': {'app': 'team-g-backend'},  # 라벨 (검색용)
            'version': '1',
        }
        _active_handlers = ['console', 'file', 'loki']
    except ImportError:
        _active_handlers = ['console', 'file']
else:
    _active_handlers = ['console', 'file']'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• ConsoleHandler: 개발 시 터미널에서 실시간 로그 확인')
    doc.add_paragraph('• RotatingFileHandler: 로그 파일이 10MB 초과시 자동 로테이션, 5개까지 보관')
    doc.add_paragraph('• LokiHandler: HTTP POST로 Loki 서버에 로그 전송')
    doc.add_paragraph('• tags: Loki에서 로그 검색시 사용하는 라벨 ({app="team-g-backend"})')
    doc.add_paragraph('')

    doc.add_heading('3.3.3 포매터 설정', level=3)

    doc.add_paragraph('# config/settings.py (라인 296-304)')
    code_block = doc.add_paragraph()
    code_block.add_run('''LOGGING = {
    'version': 1,
    'disable_existing_loggers': False,
    'formatters': {
        'verbose': {
            'format': '{levelname} {asctime} {module} {process:d} {thread:d} {message}',
            'style': '{',
        },
        'json': {
            'format': '{"timestamp": "%(asctime)s", "level": "%(levelname)s", ...}',
        },
    },
    ...
}'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• verbose: 사람이 읽기 쉬운 형식 (콘솔용)')
    doc.add_paragraph('• json: 구조화된 JSON 형식 (파일/분석용)')
    doc.add_paragraph('')

    doc.add_heading('3.3.4 로거 설정', level=3)

    doc.add_paragraph('# config/settings.py (라인 310-322)')
    code_block = doc.add_paragraph()
    code_block.add_run('''    'loggers': {
        'django': {
            'handlers': _active_handlers,  # ['console', 'file', 'loki']
            'level': os.getenv('DJANGO_LOG_LEVEL', 'INFO'),
            'propagate': False,
        },
        'celery': {
            'handlers': _active_handlers,
            'level': 'INFO',
            'propagate': False,
        },
    },'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• 각 로거(django, celery)에 모든 활성 핸들러 적용')
    doc.add_paragraph('• propagate=False: 상위 로거로 전파 방지 (중복 로그 방지)')
    doc.add_paragraph('')

    doc.add_heading('3.4 Grafana 시각화', level=2)

    doc.add_paragraph(
        'Grafana에서 Loki 데이터소스를 추가하면 Explore 탭에서 로그를 조회할 수 있습니다.'
    )
    doc.add_paragraph('')
    doc.add_paragraph('LogQL 쿼리 예시:')
    doc.add_paragraph('• {app="team-g-backend"} - 모든 로그')
    doc.add_paragraph('• {app="team-g-backend"} |= "ERROR" - 에러 로그만')
    doc.add_paragraph('• {app="team-g-backend"} | json | level="ERROR" - JSON 파싱 후 에러 필터')
    doc.add_paragraph('')

    # =========================================================================
    # 4. 메트릭 (Metrics) 구현
    # =========================================================================
    doc.add_heading('4. 메트릭 (Metrics) 구현', level=1)

    doc.add_heading('4.1 기술 스택', level=2)

    doc.add_paragraph('• prometheus_client: Python용 Prometheus 클라이언트 라이브러리')
    doc.add_paragraph('• django_prometheus: Django HTTP 메트릭 자동 수집 미들웨어')
    doc.add_paragraph('• Prometheus: 메트릭 저장소 (Pull 모델)')
    doc.add_paragraph('• Pushgateway: Celery처럼 짧게 실행되는 작업의 메트릭을 Push로 수집')
    doc.add_paragraph('• Grafana: 메트릭 시각화 및 알림')

    doc.add_heading('4.2 동작 원리', level=2)

    doc.add_paragraph('Prometheus는 두 가지 수집 모델을 지원합니다:')
    doc.add_paragraph('')

    doc.add_heading('4.2.1 Pull 모델 (Django 웹 서버)', level=3)
    doc.add_paragraph('Prometheus가 주기적으로 /metrics 엔드포인트에 HTTP GET 요청')
    doc.add_paragraph('')
    doc.add_paragraph('[Django] ← GET /metrics ← [Prometheus] (15초마다)')
    doc.add_paragraph('')

    doc.add_heading('4.2.2 Push 모델 (Celery 워커)', level=3)
    doc.add_paragraph('Celery 태스크 완료 후 Pushgateway로 메트릭을 Push')
    doc.add_paragraph('')
    doc.add_paragraph('[Celery] → POST /metrics/job/celery → [Pushgateway] ← GET ← [Prometheus]')
    doc.add_paragraph('')

    doc.add_heading('4.3 메트릭 타입', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '타입'
    header[1].text = '설명'
    header[2].text = '예시'

    row = table.rows[1].cells
    row[0].text = 'Counter'
    row[1].text = '단조 증가하는 누적값. 감소 불가.'
    row[2].text = '총 요청 수, 에러 발생 횟수'

    row = table.rows[2].cells
    row[0].text = 'Gauge'
    row[1].text = '증감 가능한 현재 값'
    row[2].text = '현재 처리 중인 작업 수, 메모리 사용량'

    row = table.rows[3].cells
    row[0].text = 'Histogram'
    row[1].text = '값의 분포를 버킷으로 측정'
    row[2].text = '응답 시간 분포 (p50, p95, p99)'

    doc.add_paragraph('')

    doc.add_heading('4.4 구현 코드 설명 (services/metrics.py)', level=2)

    doc.add_heading('4.4.1 메트릭 정의', level=3)

    doc.add_paragraph('# services/metrics.py (라인 38-60)')
    code_block = doc.add_paragraph()
    code_block.add_run('''# Counter: 분석 완료 수 (성공/실패별)
ANALYSIS_TOTAL = Counter(
    'teamg_analysis_total',        # 메트릭 이름
    'Total number of image analyses processed',  # 설명
    ['status']                     # 라벨 (success, failed)
)

# Histogram: 파이프라인 단계별 소요시간
ANALYSIS_DURATION = Histogram(
    'teamg_analysis_duration_seconds',
    'Time spent in each analysis pipeline stage',
    ['stage'],                     # 라벨 (detect_objects, extract_features, ...)
    buckets=(0.1, 0.5, 1.0, 2.0, 5.0, 10.0, 30.0, 60.0, 120.0)  # 버킷 경계
)

# Gauge: 현재 진행 중인 분석 수
ANALYSIS_IN_PROGRESS = Gauge(
    'teamg_analysis_in_progress',
    'Number of analyses currently being processed'
)'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• Counter - labels: 라벨로 세분화 (status=success vs status=failed)')
    doc.add_paragraph('• Histogram - buckets: 응답시간 분포 측정의 구간 설정')
    doc.add_paragraph('• 메트릭 이름 규칙: {namespace}_{name}_{unit} (예: teamg_analysis_duration_seconds)')
    doc.add_paragraph('')

    doc.add_heading('4.4.2 외부 API 메트릭', level=3)

    doc.add_paragraph('# services/metrics.py (라인 67-84)')
    code_block = doc.add_paragraph()
    code_block.add_run('''# 외부 API 호출 수 (서비스별, 상태별)
EXTERNAL_API_REQUESTS = Counter(
    'teamg_external_api_requests_total',
    'External API calls by service and status',
    ['service', 'status']  # service: google_vision, claude, fashn / status: success, error
)

# 외부 API 응답 시간
EXTERNAL_API_DURATION = Histogram(
    'teamg_external_api_duration_seconds',
    'External API call latency by service',
    ['service'],
    buckets=(0.1, 0.25, 0.5, 1.0, 2.0, 5.0, 10.0, 30.0)
)

# 외부 API 에러 (서비스별, 에러 타입별)
EXTERNAL_API_ERRORS = Counter(
    'teamg_external_api_errors_total',
    'External API errors by service and error type',
    ['service', 'error_type']  # error_type: timeout, rate_limit, auth, server_error
)'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')

    doc.add_heading('4.4.3 Context Manager로 메트릭 기록', level=3)

    doc.add_paragraph('# services/metrics.py (라인 135-155)')
    code_block = doc.add_paragraph()
    code_block.add_run('''@contextmanager
def record_api_call(service: str):
    """
    외부 API 호출을 감싸서 자동으로 메트릭을 기록하는 context manager.

    사용법:
        with record_api_call('google_vision'):
            response = vision_client.detect_objects(image)
    """
    start_time = time.time()
    try:
        yield
        # 성공시 success 카운터 증가
        EXTERNAL_API_REQUESTS.labels(service=service, status='success').inc()
    except Exception as e:
        # 실패시 error 카운터 증가
        EXTERNAL_API_REQUESTS.labels(service=service, status='error').inc()
        error_type = _classify_error(e)  # timeout, rate_limit, auth 등으로 분류
        EXTERNAL_API_ERRORS.labels(service=service, error_type=error_type).inc()
        raise
    finally:
        # 항상 소요시간 기록
        duration = time.time() - start_time
        EXTERNAL_API_DURATION.labels(service=service).observe(duration)'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• with 구문으로 감싸면 자동으로 시작/종료 시간, 성공/실패 기록')
    doc.add_paragraph('• _classify_error(): 예외 메시지를 분석하여 에러 타입 분류')
    doc.add_paragraph('')

    doc.add_heading('4.4.4 Pushgateway로 메트릭 전송', level=3)

    doc.add_paragraph('# services/metrics.py (라인 158-171)')
    code_block = doc.add_paragraph()
    code_block.add_run('''def push_metrics(job_name: str = 'celery_worker'):
    """
    Celery 워커에서 Pushgateway로 메트릭 푸시.

    Celery는 짧게 실행되고 종료되므로 Prometheus가 Pull할 시간이 없음.
    따라서 태스크 완료 후 Pushgateway로 직접 Push.
    """
    try:
        push_to_gateway(PUSHGATEWAY_URL, job=job_name, registry=REGISTRY)
        logger.debug(f"Metrics pushed to {PUSHGATEWAY_URL}")
    except Exception as e:
        logger.warning(f"Failed to push metrics to Pushgateway: {e}")'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')

    doc.add_heading('4.5 Django Prometheus 미들웨어', level=2)

    doc.add_paragraph('# config/settings.py (라인 49-60)')
    code_block = doc.add_paragraph()
    code_block.add_run('''MIDDLEWARE = [
    'django_prometheus.middleware.PrometheusBeforeMiddleware',  # 요청 시작
    ...
    'django_prometheus.middleware.PrometheusAfterMiddleware',   # 요청 완료
]'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• PrometheusBeforeMiddleware: 요청 시작 시간 기록')
    doc.add_paragraph('• PrometheusAfterMiddleware: 요청 완료 후 응답시간, 상태코드 등 자동 기록')
    doc.add_paragraph('• /metrics 엔드포인트 자동 생성')
    doc.add_paragraph('')

    doc.add_heading('4.6 Grafana 시각화', level=2)

    doc.add_paragraph('PromQL 쿼리 예시:')
    doc.add_paragraph('')
    doc.add_paragraph('• 분석 성공률:')
    code_block = doc.add_paragraph()
    code_block.add_run('sum(rate(teamg_analysis_total{status="success"}[5m])) / sum(rate(teamg_analysis_total[5m])) * 100').font.name = 'Courier New'
    doc.add_paragraph('')
    doc.add_paragraph('• 외부 API 평균 응답시간:')
    code_block = doc.add_paragraph()
    code_block.add_run('histogram_quantile(0.95, rate(teamg_external_api_duration_seconds_bucket[5m]))').font.name = 'Courier New'
    doc.add_paragraph('')
    doc.add_paragraph('• 현재 진행 중인 분석:')
    code_block = doc.add_paragraph()
    code_block.add_run('teamg_analysis_in_progress').font.name = 'Courier New'
    doc.add_paragraph('')

    # =========================================================================
    # 5. 트레이싱 (Tracing) 구현
    # =========================================================================
    doc.add_heading('5. 트레이싱 (Tracing) 구현', level=1)

    doc.add_heading('5.1 기술 스택', level=2)

    doc.add_paragraph('• OpenTelemetry SDK: 분산 트레이싱의 표준 API/SDK')
    doc.add_paragraph('• OpenTelemetry Instrumentation: 자동 계측 라이브러리')
    doc.add_paragraph('  - opentelemetry-instrumentation-django: Django 요청 자동 추적')
    doc.add_paragraph('  - opentelemetry-instrumentation-celery: Celery 태스크 자동 추적')
    doc.add_paragraph('  - opentelemetry-instrumentation-requests: HTTP 클라이언트 자동 추적')
    doc.add_paragraph('  - opentelemetry-instrumentation-logging: 로그에 trace_id 자동 주입')
    doc.add_paragraph('• opentelemetry-exporter-jaeger: Jaeger로 트레이스 데이터 전송')
    doc.add_paragraph('• Jaeger: 분산 트레이싱 저장소 및 UI')

    doc.add_heading('5.2 동작 원리', level=2)

    doc.add_paragraph('트레이싱의 핵심 개념:')
    doc.add_paragraph('')

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '개념'
    header[1].text = '설명'

    row = table.rows[1].cells
    row[0].text = 'Trace'
    row[1].text = '하나의 요청이 시스템을 통과하는 전체 여정'

    row = table.rows[2].cells
    row[0].text = 'Span'
    row[1].text = 'Trace 내의 개별 작업 단위 (함수 호출, API 요청 등)'

    row = table.rows[3].cells
    row[0].text = 'Trace ID'
    row[1].text = '전체 요청을 식별하는 고유 ID (32자리 16진수)'

    doc.add_paragraph('')
    doc.add_paragraph('트레이스 흐름 예시:')
    doc.add_paragraph('')
    doc.add_paragraph('Trace ID: abc123...')
    doc.add_paragraph('├─ Span: POST /api/v1/analyses (Django)')
    doc.add_paragraph('│  ├─ Span: process_image_analysis (Celery)')
    doc.add_paragraph('│  │  ├─ Span: POST vision.googleapis.com (Google Vision)')
    doc.add_paragraph('│  │  ├─ Span: POST api.anthropic.com (Claude)')
    doc.add_paragraph('│  │  └─ Span: POST opensearch:9200 (OpenSearch)')
    doc.add_paragraph('')

    doc.add_heading('5.3 구현 코드 설명 (config/tracing.py)', level=2)

    doc.add_heading('5.3.1 초기화 함수', level=3)

    doc.add_paragraph('# config/tracing.py (라인 21-43)')
    code_block = doc.add_paragraph()
    code_block.add_run('''def init_tracing(service_name: str = "team-g-backend"):
    """
    OpenTelemetry 트레이싱 초기화.
    애플리케이션 시작 시 한 번만 호출됨.
    """
    global _tracing_initialized

    if _tracing_initialized:
        logger.debug("Tracing already initialized, skipping")
        return

    # 환경변수로 활성화/비활성화 제어
    tracing_enabled = os.getenv('TRACING_ENABLED', 'true').lower() == 'true'
    if not tracing_enabled:
        logger.info("Tracing is disabled via TRACING_ENABLED=false")
        return'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• _tracing_initialized: 중복 초기화 방지 플래그')
    doc.add_paragraph('• TRACING_ENABLED: 환경변수로 트레이싱 on/off 제어')
    doc.add_paragraph('')

    doc.add_heading('5.3.2 Resource 및 Exporter 설정', level=3)

    doc.add_paragraph('# config/tracing.py (라인 55-68)')
    code_block = doc.add_paragraph()
    code_block.add_run('''        # Resource: 서비스 메타데이터
        resource = Resource.create({
            "service.name": service_name,      # Jaeger UI에 표시될 서비스 이름
            "service.version": "1.0.0",
        })

        # Jaeger Exporter 설정
        jaeger_host = os.getenv('JAEGER_HOST', 'localhost')
        jaeger_port = int(os.getenv('JAEGER_PORT', '6831'))

        jaeger_exporter = JaegerExporter(
            agent_host_name=jaeger_host,
            agent_port=jaeger_port,           # UDP 6831 포트
        )'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• Resource: 이 서비스의 메타데이터 (이름, 버전)')
    doc.add_paragraph('• JaegerExporter: UDP 프로토콜로 Jaeger Agent에 span 전송')
    doc.add_paragraph('')

    doc.add_heading('5.3.3 TracerProvider 설정', level=3)

    doc.add_paragraph('# config/tracing.py (라인 70-73)')
    code_block = doc.add_paragraph()
    code_block.add_run('''        # TracerProvider: 트레이서를 생성하고 관리
        provider = TracerProvider(resource=resource)

        # BatchSpanProcessor: span을 모아서 일괄 전송 (성능 최적화)
        provider.add_span_processor(BatchSpanProcessor(jaeger_exporter))

        # 전역 TracerProvider로 등록
        trace.set_tracer_provider(provider)'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• TracerProvider: span을 생성하는 팩토리')
    doc.add_paragraph('• BatchSpanProcessor: 매 span마다 전송하지 않고 배치로 모아서 전송')
    doc.add_paragraph('')

    doc.add_heading('5.3.4 자동 계측 (Auto-Instrumentation)', level=3)

    doc.add_paragraph('# config/tracing.py (라인 75-85)')
    code_block = doc.add_paragraph()
    code_block.add_run('''        # Django: 모든 HTTP 요청/응답 자동 추적
        DjangoInstrumentor().instrument()

        # Celery: 모든 태스크 실행 자동 추적
        CeleryInstrumentor().instrument()

        # requests 라이브러리: 외부 API 호출 자동 추적
        RequestsInstrumentor().instrument()

        # 로그에 trace_id 자동 주입
        LoggingInstrumentor().instrument(set_logging_format=True)'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• DjangoInstrumentor: Django 뷰 함수 진입/종료 시 자동으로 span 생성')
    doc.add_paragraph('• CeleryInstrumentor: Celery 태스크 실행 시 자동으로 span 생성')
    doc.add_paragraph('• RequestsInstrumentor: requests.get/post 호출 시 자동으로 span 생성')
    doc.add_paragraph('• LoggingInstrumentor: 모든 로그에 현재 trace_id를 자동 포함')
    doc.add_paragraph('')

    doc.add_heading('5.4 초기화 시점', level=2)

    doc.add_heading('5.4.1 Django 웹 서버 (config/wsgi.py)', level=3)

    doc.add_paragraph('# config/wsgi.py (라인 16-22)')
    code_block = doc.add_paragraph()
    code_block.add_run('''# Django 애플리케이션 로드 전에 트레이싱 초기화
try:
    from config.tracing import init_tracing
    init_tracing(service_name="team-g-backend")  # 서비스 이름 지정
except ImportError:
    pass  # 트레이싱 패키지 미설치 시 무시

application = get_wsgi_application()'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')

    doc.add_heading('5.4.2 Celery 워커 (config/celery.py)', level=3)

    doc.add_paragraph('# config/celery.py (라인 21-28)')
    code_block = doc.add_paragraph()
    code_block.add_run('''@worker_process_init.connect
def init_worker_tracing(**kwargs):
    """
    Celery 워커 프로세스 시작 시 트레이싱 초기화.
    각 워커 프로세스마다 한 번씩 호출됨.
    """
    try:
        from config.tracing import init_tracing
        init_tracing(service_name="team-g-celery-worker")  # 다른 서비스 이름
    except ImportError:
        pass'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')
    doc.add_paragraph('설명:')
    doc.add_paragraph('• @worker_process_init.connect: Celery 워커 시작 시그널에 연결')
    doc.add_paragraph('• 서비스 이름을 다르게 지정하여 Jaeger UI에서 구분')
    doc.add_paragraph('')

    doc.add_heading('5.5 Trace ID 조회 유틸리티', level=2)

    doc.add_paragraph('# config/tracing.py (라인 96-112)')
    code_block = doc.add_paragraph()
    code_block.add_run('''def get_current_trace_id() -> str:
    """
    현재 실행 중인 span의 trace_id를 반환.
    로그에 수동으로 trace_id를 추가할 때 사용.
    """
    try:
        from opentelemetry import trace
        span = trace.get_current_span()
        if span:
            trace_id = span.get_span_context().trace_id
            if trace_id:
                return format(trace_id, '032x')  # 32자리 16진수로 포맷
    except Exception:
        pass
    return ""'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')

    doc.add_heading('5.6 Grafana/Jaeger 시각화', level=2)

    doc.add_paragraph('Jaeger UI (http://localhost:16686):')
    doc.add_paragraph('• Service 드롭다운에서 "team-g-backend" 또는 "team-g-celery-worker" 선택')
    doc.add_paragraph('• Find Traces 클릭하여 트레이스 목록 조회')
    doc.add_paragraph('• 트레이스 클릭 시 상세 span 타임라인 확인')
    doc.add_paragraph('')
    doc.add_paragraph('Grafana에서 Jaeger 연동:')
    doc.add_paragraph('• Data Sources에 Jaeger 추가 (http://localhost:16686)')
    doc.add_paragraph('• Explore에서 트레이스 검색 가능')
    doc.add_paragraph('')

    # =========================================================================
    # 6. 세 가지 데이터의 연결
    # =========================================================================
    doc.add_heading('6. 세 가지 데이터의 연결 (Trace ID)', level=1)

    doc.add_paragraph(
        '로그, 메트릭, 트레이스는 Trace ID를 통해 연결됩니다. '
        'Trace ID는 하나의 요청을 식별하는 고유 ID로, 세 데이터 모두에서 동일하게 사용됩니다.'
    )

    doc.add_heading('6.1 연결 흐름', level=2)

    doc.add_paragraph('1. 사용자가 POST /api/v1/analyses 요청')
    doc.add_paragraph('2. DjangoInstrumentor가 자동으로 Trace ID 생성 (예: abc123...)')
    doc.add_paragraph('3. 해당 요청 처리 중 발생하는 모든 로그에 trace_id 자동 포함')
    doc.add_paragraph('4. Celery 태스크로 전달될 때 trace_id도 함께 전파')
    doc.add_paragraph('5. 외부 API 호출 시에도 동일한 trace_id로 span 생성')
    doc.add_paragraph('')

    doc.add_heading('6.2 실제 활용 예시', level=2)

    doc.add_paragraph('문제 상황: "분석 API가 간헐적으로 느림"')
    doc.add_paragraph('')
    doc.add_paragraph('해결 과정:')
    doc.add_paragraph('1. Grafana Metrics에서 응답시간 급증 구간 확인')
    doc.add_paragraph('2. 해당 시간대의 Logs에서 에러/경고 로그 검색')
    doc.add_paragraph('3. 로그에서 trace_id 추출 (예: "trace_id": "abc123...")')
    doc.add_paragraph('4. Jaeger에서 해당 trace_id로 검색')
    doc.add_paragraph('5. 트레이스에서 어느 span(외부 API 호출)이 오래 걸렸는지 확인')
    doc.add_paragraph('6. 근본 원인 파악 (예: Google Vision API 타임아웃)')
    doc.add_paragraph('')

    doc.add_heading('6.3 LoggingInstrumentor의 역할', level=2)

    code_block = doc.add_paragraph()
    code_block.add_run('''# 자동 계측 전 로그
{"level": "INFO", "message": "Analysis started"}

# 자동 계측 후 로그 (trace_id 자동 주입)
{"level": "INFO", "message": "Analysis started",
 "otelTraceID": "abc123...", "otelSpanID": "def456..."}'''
    ).font.name = 'Courier New'

    doc.add_paragraph('')

    # =========================================================================
    # 7. 환경 변수 요약
    # =========================================================================
    doc.add_heading('7. 환경 변수 요약', level=1)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '변수명'
    header[1].text = '기본값'
    header[2].text = '설명'

    env_vars = [
        ('LOKI_URL', 'http://localhost:3100/loki/api/v1/push', 'Loki push API 엔드포인트'),
        ('LOKI_ENABLED', 'true', 'Loki 핸들러 활성화 여부'),
        ('JAEGER_HOST', 'localhost', 'Jaeger Agent 호스트'),
        ('JAEGER_PORT', '6831', 'Jaeger Agent UDP 포트'),
        ('TRACING_ENABLED', 'true', 'OpenTelemetry 트레이싱 활성화'),
        ('PUSHGATEWAY_URL', 'localhost:9091', 'Prometheus Pushgateway 주소'),
    ]

    for i, (name, default, desc) in enumerate(env_vars, 1):
        row = table.rows[i].cells
        row[0].text = name
        row[1].text = default
        row[2].text = desc

    doc.add_paragraph('')

    # =========================================================================
    # 8. 파일 목록 요약
    # =========================================================================
    doc.add_heading('8. 파일 목록 요약', level=1)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '파일'
    header[1].text = '용도'

    files = [
        ('config/tracing.py', 'OpenTelemetry 초기화 코드'),
        ('config/settings.py', 'Loki 핸들러 설정, 로깅 설정'),
        ('config/celery.py', 'Celery 워커 트레이싱 초기화'),
        ('config/wsgi.py', 'Django 웹 서버 트레이싱 초기화'),
        ('services/metrics.py', 'Prometheus 커스텀 메트릭 정의'),
        ('.env', 'LOKI_URL, JAEGER_HOST 등 환경변수'),
    ]

    for i, (file, desc) in enumerate(files, 1):
        row = table.rows[i].cells
        row[0].text = file
        row[1].text = desc

    doc.add_paragraph('')

    # 저장
    doc.save('/Users/ijeong/Desktop/Team_G/Team_G_Backend/Team_G_Observability_구현문서.docx')
    print("문서 생성 완료: Team_G_Observability_구현문서.docx")


if __name__ == '__main__':
    create_document()
