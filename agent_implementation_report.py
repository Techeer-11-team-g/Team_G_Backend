"""
AI 패션 어시스턴트 에이전트 구현 보고서 생성 스크립트
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def create_report():
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # 제목
    title = doc.add_heading('AI 패션 어시스턴트 에이전트 구현 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 메타 정보
    meta = doc.add_paragraph()
    meta.add_run(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}\n').bold = True
    meta.add_run('프로젝트: Team_G Backend\n')
    meta.add_run('버전: 1.0.0')
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # =========================================================================
    # 1. 개요
    # =========================================================================
    doc.add_heading('1. 개요', level=1)

    doc.add_paragraph(
        'AI 패션 어시스턴트 에이전트 시스템을 Team_G Backend에 구현 완료했습니다. '
        '이 시스템은 기존의 RESTful API 방식을 대체하여 자연어 기반의 통합 채팅 인터페이스를 제공합니다.'
    )

    doc.add_heading('1.1 구현 범위', level=2)
    items = [
        'agents Django 앱 생성 및 구조화',
        'MainOrchestrator (의도 분류 및 라우팅)',
        'SearchAgent (이미지/텍스트 검색)',
        'FittingAgent (가상 피팅)',
        'CommerceAgent (장바구니/주문)',
        'REST API 엔드포인트 및 직렬화',
        'Redis 기반 세션 관리',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================================
    # 2. 변경사항
    # =========================================================================
    doc.add_heading('2. 변경사항', level=1)

    doc.add_heading('2.1 신규 파일', level=2)

    # 테이블 생성
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '파일 경로'
    hdr_cells[1].text = '역할'
    hdr_cells[2].text = '라인 수'

    new_files = [
        ('agents/schemas.py', '의도 분류 스키마 및 상수 정의', '~150'),
        ('agents/response_builder.py', '응답 포맷 빌더 클래스', '~200'),
        ('agents/orchestrator.py', '메인 오케스트레이터', '~350'),
        ('agents/sub_agents/__init__.py', '서브에이전트 패키지', '~10'),
        ('agents/sub_agents/search_agent.py', '검색 에이전트', '~250'),
        ('agents/sub_agents/fitting_agent.py', '피팅 에이전트', '~200'),
        ('agents/sub_agents/commerce_agent.py', '커머스 에이전트', '~300'),
        ('agents/views.py', 'API 뷰 (ChatView 등)', '~230'),
        ('agents/serializers.py', '요청/응답 직렬화', '~90'),
        ('agents/urls.py', 'URL 라우팅', '~20'),
    ]

    for path, role, lines in new_files:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = role
        row_cells[2].text = lines

    doc.add_paragraph()

    doc.add_heading('2.2 수정된 파일', level=2)

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '파일 경로'
    hdr_cells2[1].text = '변경 내용'
    hdr_cells2[2].text = '영향도'

    modified_files = [
        ('config/settings.py', "INSTALLED_APPS에 'agents' 추가, agents 로거 추가", '최소'),
        ('config/urls.py', "agents.urls include 추가", '최소'),
    ]

    for path, change, impact in modified_files:
        row_cells = table2.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = change
        row_cells[2].text = impact

    doc.add_paragraph()

    doc.add_heading('2.3 새로운 API 엔드포인트', level=2)

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = '메서드'
    hdr_cells3[1].text = '경로'
    hdr_cells3[2].text = '설명'
    hdr_cells3[3].text = '인증'

    endpoints = [
        ('POST', '/api/v1/chat', '채팅 메시지 처리 (통합 인터페이스)', '필수'),
        ('POST', '/api/v1/chat/status', '분석/피팅 상태 확인', '필수'),
        ('GET', '/api/v1/chat/sessions', '세션 목록 조회', '필수'),
        ('GET', '/api/v1/chat/sessions/<id>', '특정 세션 조회', '필수'),
        ('DELETE', '/api/v1/chat/sessions/<id>', '세션 삭제', '필수'),
    ]

    for method, path, desc, auth in endpoints:
        row_cells = table3.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = path
        row_cells[2].text = desc
        row_cells[3].text = auth

    # =========================================================================
    # 3. 아키텍처 상세
    # =========================================================================
    doc.add_heading('3. 아키텍처 상세', level=1)

    doc.add_heading('3.1 에이전트 구조', level=2)
    doc.add_paragraph(
        '에이전트 시스템은 계층적 구조로 설계되었습니다:'
    )

    structure_text = """
┌─────────────────────────────────────────┐
│              ChatView                    │
│         (REST API 진입점)               │
└─────────────────────────────────────────┘
                    │
                    ▼
┌─────────────────────────────────────────┐
│          MainOrchestrator               │
│     • 의도 분류 (LangChain + 키워드)     │
│     • 컨텍스트 관리 (Redis)              │
│     • 서브에이전트 라우팅                │
└─────────────────────────────────────────┘
          │           │           │
          ▼           ▼           ▼
    ┌──────────┐ ┌──────────┐ ┌──────────┐
    │  Search  │ │ Fitting  │ │ Commerce │
    │  Agent   │ │  Agent   │ │  Agent   │
    └──────────┘ └──────────┘ └──────────┘
          │           │           │
          ▼           ▼           ▼
    ┌─────────────────────────────────────┐
    │         기존 서비스 레이어           │
    │  • LangChainService                 │
    │  • parse_refine_query_task          │
    │  • process_fitting_task             │
    │  • CartItem/Order 모델              │
    └─────────────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(structure_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('3.2 의도 분류 시스템', level=2)
    doc.add_paragraph(
        '사용자 메시지의 의도를 분류하는 하이브리드 방식을 채택했습니다:'
    )

    doc.add_paragraph('1차: LangChain Function Calling', style='List Number')
    doc.add_paragraph(
        '   - Claude/GPT 모델의 Function Calling 기능 활용\n'
        '   - INTENT_CLASSIFICATION_SCHEMA 기반 구조화된 출력\n'
        '   - 복합 의도(compound) 감지 지원'
    )

    doc.add_paragraph('2차: 키워드 기반 폴백', style='List Number')
    doc.add_paragraph(
        '   - LLM 호출 실패 시 키워드 매칭으로 폴백\n'
        '   - INTENT_KEYWORDS 딕셔너리 기반\n'
        '   - 이미지 존재 시 자동으로 search 의도로 분류'
    )

    doc.add_heading('3.3 세션 관리', level=2)
    doc.add_paragraph('Redis를 활용한 세션 관리:')

    session_items = [
        '세션 TTL: 2시간 (기존 TTL_CONVERSATION과 동일)',
        '키 패턴: agent:session:{session_id}',
        '대화 턴: agent:session:{session_id}:turns (List 타입)',
        '저장 데이터: user_id, current_analysis_id, search_results, user_image_url 등',
    ]
    for item in session_items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================================
    # 4. 기존 코드와의 차이점
    # =========================================================================
    doc.add_heading('4. 기존 코드와의 차이점', level=1)

    doc.add_heading('4.1 API 설계 패러다임 변화', level=2)

    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    hdr_cells4 = table4.rows[0].cells
    hdr_cells4[0].text = '항목'
    hdr_cells4[1].text = '기존 방식'
    hdr_cells4[2].text = '에이전트 방식'

    comparisons = [
        ('API 구조', '기능별 개별 엔드포인트\n(POST /analyses, POST /fitting-images 등)', '통합 채팅 엔드포인트\n(POST /api/v1/chat)'),
        ('사용자 입력', 'JSON 필드 명시적 지정\n(query, category, brand 등)', '자연어 메시지\n("빨간색 나이키 운동화 보여줘")'),
        ('상태 관리', '클라이언트에서 상태 관리\n(analysis_id 저장 필요)', '서버 세션 기반\n(Redis에서 컨텍스트 유지)'),
        ('흐름 제어', '클라이언트가 API 순서 결정\n(분석 → 피팅 → 주문)', '에이전트가 자동 판단 및 라우팅'),
        ('에러 처리', 'HTTP 상태 코드 기반', '자연어 에러 메시지 + suggestions'),
    ]

    for item, old, new in comparisons:
        row_cells = table4.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = old
        row_cells[2].text = new

    doc.add_paragraph()

    doc.add_heading('4.2 코드 재사용성', level=2)
    doc.add_paragraph(
        '에이전트 시스템은 기존 서비스 레이어를 최대한 재사용합니다:'
    )

    reuse_items = [
        ('검색 기능', 'parse_refine_query_task (analyses/tasks/refine.py) 그대로 활용'),
        ('이미지 분석', 'process_image_analysis 태스크 재사용'),
        ('가상 피팅', 'process_fitting_task (fittings/tasks.py) 재사용'),
        ('상품 데이터', 'Product, SizeCode 모델 그대로 사용'),
        ('주문 시스템', 'CartItem, Order, OrderItem 모델 재사용'),
        ('LLM 서비스', 'LangChainService 싱글톤 패턴 유지'),
        ('Redis 서비스', 'get_redis_service() 팩토리 함수 사용'),
    ]

    table5 = doc.add_table(rows=1, cols=2)
    table5.style = 'Table Grid'
    hdr_cells5 = table5.rows[0].cells
    hdr_cells5[0].text = '기능'
    hdr_cells5[1].text = '재사용 방식'

    for func, reuse in reuse_items:
        row_cells = table5.add_row().cells
        row_cells[0].text = func
        row_cells[1].text = reuse

    doc.add_paragraph()

    doc.add_heading('4.3 기존 API와의 공존', level=2)
    doc.add_paragraph(
        '에이전트 API는 기존 RESTful API와 완전히 독립적으로 동작합니다:'
    )

    coexist_items = [
        '기존 엔드포인트 변경 없음 (/api/v1/analyses, /api/v1/fitting-images 등)',
        '기존 클라이언트 코드 수정 불필요',
        '점진적 마이그레이션 가능 (새 기능은 에이전트로, 기존 기능은 그대로)',
        '동일 백엔드 인프라 공유 (DB, Redis, Celery, OpenSearch)',
    ]
    for item in coexist_items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================================
    # 5. 기대 효과
    # =========================================================================
    doc.add_heading('5. 기대 효과', level=1)

    doc.add_heading('5.1 사용자 경험 향상', level=2)

    ux_effects = [
        ('학습 곡선 감소', '복잡한 API 파라미터 대신 자연어로 상호작용'),
        ('컨텍스트 유지', '이전 대화 내용을 기반으로 연속적인 대화 가능'),
        ('스마트 제안', '다음 액션 제안으로 사용자 가이드 (suggestions 필드)'),
        ('에러 친화성', '기술적 에러 대신 이해하기 쉬운 메시지 제공'),
        ('원스톱 쇼핑', '검색 → 피팅 → 주문까지 하나의 채팅창에서 완료'),
    ]

    for title, desc in ux_effects:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_heading('5.2 개발/운영 효율성', level=2)

    dev_effects = [
        ('프론트엔드 단순화', '여러 API 호출 로직 대신 단일 채팅 인터페이스'),
        ('유지보수 용이', '의도별 모듈화로 독립적 수정/테스트 가능'),
        ('확장성', '새 서브에이전트 추가로 기능 확장 용이'),
        ('모니터링', '기존 로깅/트레이싱 인프라 활용'),
        ('A/B 테스트', '기존 API와 에이전트 API 병행 운영으로 비교 가능'),
    ]

    for title, desc in dev_effects:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_heading('5.3 비즈니스 가치', level=2)

    biz_effects = [
        ('전환율 향상', '자연스러운 쇼핑 경험으로 이탈률 감소 예상'),
        ('평균 주문 금액', '피팅 체험 → 구매 연결 촉진'),
        ('재방문율', '개인화된 대화 경험으로 충성도 향상'),
        ('차별화', 'AI 쇼핑 어시스턴트로 경쟁 서비스 대비 혁신성'),
    ]

    for title, desc in biz_effects:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # =========================================================================
    # 6. 사용 예시
    # =========================================================================
    doc.add_heading('6. 사용 예시', level=1)

    doc.add_heading('6.1 검색 시나리오', level=2)

    example1 = """
# 요청
POST /api/v1/chat
{
    "message": "나이키 에어맥스 흰색 260 사이즈 있는 거 보여줘",
    "session_id": null  // 새 세션 시작
}

# 응답
{
    "session_id": "abc123...",
    "response": {
        "text": "나이키 에어맥스 흰색 260 사이즈 상품을 찾았어요!",
        "type": "search_results",
        "data": {
            "products": [...],
            "total_count": 15
        },
        "suggestions": [
            {"label": "첫 번째 상품 입어보기", "action": "fit:1"},
            {"label": "더 많은 상품 보기", "action": "more"},
            {"label": "가격순 정렬", "action": "sort:price"}
        ]
    }
}
"""
    p = doc.add_paragraph()
    run = p.add_run(example1)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('6.2 피팅 시나리오', level=2)

    example2 = """
# 요청 (이전 대화 컨텍스트 유지)
POST /api/v1/chat
{
    "message": "2번 상품 입어볼래",
    "session_id": "abc123..."
}

# 응답
{
    "session_id": "abc123...",
    "response": {
        "text": "가상 피팅을 시작합니다. 잠시 후 결과를 확인해 주세요.",
        "type": "fitting_started",
        "data": {
            "fitting_id": 456,
            "status_url": "/api/v1/chat/status"
        },
        "suggestions": [
            {"label": "상태 확인", "action": "status:fitting:456"}
        ]
    }
}
"""
    p = doc.add_paragraph()
    run = p.add_run(example2)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('6.3 구매 시나리오', level=2)

    example3 = """
# 요청
POST /api/v1/chat
{
    "message": "이거 265 사이즈로 장바구니에 담아줘",
    "session_id": "abc123..."
}

# 응답
{
    "session_id": "abc123...",
    "response": {
        "text": "Nike Air Max 90 (265)을(를) 장바구니에 담았습니다.",
        "type": "cart_added",
        "data": {
            "product_id": 789,
            "size": "265",
            "quantity": 1
        },
        "suggestions": [
            {"label": "장바구니 보기", "action": "view_cart"},
            {"label": "바로 구매하기", "action": "checkout"},
            {"label": "쇼핑 계속하기", "action": "continue"}
        ]
    }
}
"""
    p = doc.add_paragraph()
    run = p.add_run(example3)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # =========================================================================
    # 7. 향후 계획
    # =========================================================================
    doc.add_heading('7. 향후 계획', level=1)

    doc.add_heading('7.1 1단계: 안정화 (1-2주)', level=2)
    phase1_items = [
        '통합 테스트 및 엣지 케이스 처리',
        '에러 핸들링 강화',
        '프론트엔드 연동 테스트',
        '성능 모니터링 및 최적화',
    ]
    for item in phase1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 2단계: 기능 확장', level=2)
    phase2_items = [
        'StyleAgent 추가 (스타일 추천)',
        '복합 의도 처리 고도화',
        '대화 히스토리 영구 저장 (DB 연동)',
        '푸시 알림 연동 (피팅 완료 등)',
    ]
    for item in phase2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.3 3단계: 고도화', level=2)
    phase3_items = [
        '멀티모달 강화 (이미지 + 텍스트 동시 처리)',
        '개인화 추천 알고리즘 연동',
        'A/B 테스트 프레임워크 구축',
        '실시간 스트리밍 응답 (SSE)',
    ]
    for item in phase3_items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================================
    # 8. 결론
    # =========================================================================
    doc.add_heading('8. 결론', level=1)

    doc.add_paragraph(
        'AI 패션 어시스턴트 에이전트 시스템의 핵심 구현이 완료되었습니다. '
        '이 시스템은 기존 Team_G Backend의 서비스 레이어를 최대한 재활용하면서, '
        '사용자에게는 자연어 기반의 직관적인 쇼핑 경험을 제공합니다.'
    )

    doc.add_paragraph(
        '주요 성과:'
    )

    achievements = [
        '기존 코드 최소 수정 (settings.py, urls.py 2개 파일만 변경)',
        '신규 코드 약 1,500라인 추가 (agents 앱)',
        '기존 RESTful API와 완전 호환 (독립적 운영)',
        '확장 가능한 모듈화 구조 (새 에이전트 추가 용이)',
    ]
    for item in achievements:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '향후 통합 테스트와 프론트엔드 연동을 거쳐 실 서비스에 적용할 수 있습니다.'
    )

    # =========================================================================
    # 부록: 파일 구조
    # =========================================================================
    doc.add_heading('부록: 에이전트 앱 파일 구조', level=1)

    structure = """
agents/
├── __init__.py
├── apps.py                    # Django 앱 설정
├── urls.py                    # URL 라우팅
├── views.py                   # API 뷰 (ChatView 등)
├── serializers.py             # 요청/응답 직렬화
├── schemas.py                 # 의도 분류 스키마
├── response_builder.py        # 응답 포맷 빌더
├── orchestrator.py            # 메인 오케스트레이터
├── sub_agents/
│   ├── __init__.py
│   ├── search_agent.py        # 검색 에이전트
│   ├── fitting_agent.py       # 피팅 에이전트
│   └── commerce_agent.py      # 커머스 에이전트
└── tests/
    └── __init__.py
"""
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # 저장
    output_path = '/Users/ijeong/Desktop/테커/AI_패션_어시스턴트_구현_보고서.docx'
    doc.save(output_path)
    print(f'보고서가 생성되었습니다: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()
