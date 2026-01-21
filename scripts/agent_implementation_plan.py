#!/usr/bin/env python3
"""
AI 패션 어시스턴트 에이전트 - 상세 구현 계획서
기존 Team_G Backend 프로젝트에 에이전트 레이어 통합
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_table(doc, title, data, headers, level=3):
    if title:
        doc.add_heading(title, level=level)
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    doc.add_paragraph()


def add_code_block(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def create_implementation_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)

    # ================================================================
    # 표지
    # ================================================================
    title = doc.add_heading('AI 패션 어시스턴트', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('상세 구현 계획서')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Team_G Backend 프로젝트 통합 가이드\n').bold = True
    info.add_run('기존 인프라 활용 + 에이전트 레이어 추가\n')
    info.add_run('작성일: 2026-01-20')

    doc.add_page_break()

    # ================================================================
    # 목차
    # ================================================================
    doc.add_heading('목차', level=1)
    toc = '''
1. 프로젝트 현황 분석
   1.1 기존 아키텍처 요약
   1.2 활용 가능한 기술 스택
   1.3 기존 API 엔드포인트

2. 에이전트 통합 설계
   2.1 통합 아키텍처
   2.2 신규 컴포넌트
   2.3 기존 서비스 연동

3. 데이터 모델 확장
   3.1 신규 모델
   3.2 기존 모델 활용

4. API 엔드포인트 설계
   4.1 채팅 API
   4.2 세션 관리 API

5. 에이전트 핵심 로직
   5.1 Intent Classification
   5.2 서브 에이전트 라우팅
   5.3 응답 생성

6. 기존 서비스 연동 상세
   6.1 Search Agent ↔ 기존 분석 파이프라인
   6.2 Fitting Agent ↔ 기존 피팅 시스템
   6.3 Commerce Agent ↔ 기존 주문 시스템

7. Redis 상태 관리
   7.1 세션 스키마
   7.2 TTL 정책

8. Edge Case 처리
   8.1 입력 처리
   8.2 대화 흐름
   8.3 오류 복구

9. 구현 단계별 계획
   9.1 Phase 1: 기반 구축
   9.2 Phase 2: 검색 에이전트
   9.3 Phase 3: 피팅 에이전트
   9.4 Phase 4: 커머스 에이전트
   9.5 Phase 5: 통합 및 최적화

10. 파일 구조
'''
    doc.add_paragraph(toc)
    doc.add_page_break()

    # ================================================================
    # 1. 프로젝트 현황 분석
    # ================================================================
    doc.add_heading('1. 프로젝트 현황 분석', level=1)

    doc.add_heading('1.1 기존 아키텍처 요약', level=2)

    current_arch = '''
┌─────────────────────────────────────────────────────────────────────┐
│                       현재 Team_G Backend                           │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  [Frontend]                                                         │
│      │                                                              │
│      ▼                                                              │
│  [Django REST API]  /api/v1/                                        │
│      │                                                              │
│      ├── /uploaded-images    → 이미지 업로드                        │
│      ├── /analyses           → 이미지 분석 (POST, GET, PATCH)       │
│      ├── /fitting-images     → 가상 피팅                            │
│      ├── /cart-items         → 장바구니 CRUD                        │
│      ├── /orders             → 주문 관리                            │
│      └── /users              → 사용자 관리                          │
│                                                                     │
│  [Services Layer] (Singleton Pattern)                               │
│      ├── VisionService       → Google Vision API                    │
│      ├── EmbeddingService    → FashionCLIP (512-dim)                │
│      ├── OpenSearchService   → k-NN 벡터 검색                       │
│      ├── LangChainService    → GPT/Claude Function Calling          │
│      ├── GPT4VService        → Claude Vision 속성 추출              │
│      ├── FashnService        → The New Black 피팅 API               │
│      └── RedisService        → 상태/캐시 관리                       │
│                                                                     │
│  [Celery Tasks] (3 Queues: default, analysis, fitting)              │
│      ├── process_image_analysis  → 10단계 분석 파이프라인           │
│      ├── parse_refine_query      → 자연어 재분석                    │
│      └── process_fitting_task    → 가상 피팅 처리                   │
│                                                                     │
│  [Database]                                                         │
│      ├── MySQL: 영구 저장 (User, Product, Order, Analysis...)       │
│      └── Redis: 상태 캐시 (TTL 30분 폴링, 2시간 대화)               │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, current_arch)

    doc.add_heading('1.2 활용 가능한 기술 스택', level=2)

    available_stack = [
        ('LangChainService', 'GPT-4o-mini / Claude', '이미 구현됨', '대화 이력 관리, Function Calling'),
        ('RedisService', 'Redis', '이미 구현됨', 'TTL_CONVERSATION=2시간'),
        ('OpenSearchService', 'OpenSearch k-NN', '이미 구현됨', '하이브리드 검색'),
        ('EmbeddingService', 'FashionCLIP', '이미 구현됨', '512차원 임베딩'),
        ('VisionService', 'Google Vision', '이미 구현됨', '객체 감지'),
        ('GPT4VService', 'Claude Vision', '이미 구현됨', '속성 추출, 리랭킹'),
        ('FashnService', 'The New Black', '이미 구현됨', '가상 피팅'),
        ('parse_refine_query', 'Celery Task', '이미 구현됨', '자연어 필터 파싱'),
    ]
    add_table(doc, '', available_stack, ['서비스', '기술', '상태', '활용 포인트'])

    doc.add_heading('1.3 기존 API 엔드포인트', level=2)

    existing_apis = [
        ('POST /uploaded-images', '이미지 업로드', '에이전트가 직접 호출'),
        ('POST /analyses', '분석 시작', '에이전트가 직접 호출'),
        ('GET /analyses/{id}/status', '상태 폴링', '에이전트가 폴링'),
        ('PATCH /analyses', '자연어 재분석', '핵심 활용 포인트'),
        ('POST /user-images', '전신 이미지 등록', '피팅 전제조건'),
        ('POST /fitting-images', '피팅 요청', '에이전트가 직접 호출'),
        ('GET /cart-items', '장바구니 조회', '에이전트가 직접 호출'),
        ('POST /cart-items', '장바구니 추가', '에이전트가 직접 호출'),
        ('POST /orders', '주문 생성', '에이전트가 직접 호출'),
    ]
    add_table(doc, '', existing_apis, ['엔드포인트', '기능', '에이전트 활용'])

    doc.add_page_break()

    # ================================================================
    # 2. 에이전트 통합 설계
    # ================================================================
    doc.add_heading('2. 에이전트 통합 설계', level=1)

    doc.add_heading('2.1 통합 아키텍처', level=2)

    integrated_arch = '''
┌─────────────────────────────────────────────────────────────────────┐
│                    AI 패션 어시스턴트 통합 아키텍처                   │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  [Frontend/Client]                                                  │
│      │                                                              │
│      ▼                                                              │
│  ┌───────────────────────────────────────────────────────────────┐  │
│  │                    🧠 Agent API Layer (신규)                   │  │
│  │                                                               │  │
│  │   POST /api/v1/chat                                           │  │
│  │   ├── 메시지 + 이미지 수신                                     │  │
│  │   ├── 세션 컨텍스트 로드 (Redis)                               │  │
│  │   ├── Intent Classification (LangChainService)                 │  │
│  │   ├── 서브 에이전트 라우팅                                     │  │
│  │   ├── 응답 생성                                                │  │
│  │   └── 세션 업데이트 (Redis)                                    │  │
│  │                                                               │  │
│  └───────────────────────────┬───────────────────────────────────┘  │
│                              │                                      │
│         ┌────────────────────┼────────────────────┐                 │
│         ▼                    ▼                    ▼                 │
│   ┌───────────┐        ┌───────────┐        ┌───────────┐          │
│   │  Search   │        │  Fitting  │        │ Commerce  │          │
│   │  Agent    │        │  Agent    │        │  Agent    │          │
│   └─────┬─────┘        └─────┬─────┘        └─────┬─────┘          │
│         │                    │                    │                 │
│         ▼                    ▼                    ▼                 │
│  ┌─────────────────────────────────────────────────────────────┐   │
│  │                  기존 API / Services (활용)                  │   │
│  │                                                             │   │
│  │   POST /uploaded-images  →  VisionService                   │   │
│  │   POST /analyses         →  process_image_analysis          │   │
│  │   PATCH /analyses        →  parse_refine_query (핵심!)      │   │
│  │   POST /fitting-images   →  process_fitting_task            │   │
│  │   POST /cart-items       →  CartItem 모델                   │   │
│  │   POST /orders           →  Order 모델                      │   │
│  │                                                             │   │
│  └─────────────────────────────────────────────────────────────┘   │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, integrated_arch)

    doc.add_heading('2.2 신규 컴포넌트', level=2)

    new_components = [
        ('agents/', 'Django App', '에이전트 핵심 로직'),
        ('agents/views.py', 'ChatView', '채팅 API 엔드포인트'),
        ('agents/orchestrator.py', 'MainOrchestrator', '의도 분류, 라우팅'),
        ('agents/sub_agents/', 'SearchAgent, FittingAgent, CommerceAgent', '서브 에이전트'),
        ('agents/models.py', 'ChatSession, ChatMessage', '대화 이력 (선택)'),
        ('agents/serializers.py', 'ChatRequestSerializer, ChatResponseSerializer', '직렬화'),
        ('services/agent_service.py', 'AgentService', '에이전트 서비스 레이어'),
    ]
    add_table(doc, '', new_components, ['경로', '컴포넌트', '역할'])

    doc.add_heading('2.3 기존 서비스 연동', level=2)

    doc.add_paragraph(
        '에이전트는 기존 서비스를 최대한 재사용합니다. '
        '특히 LangChainService의 대화 이력 관리와 parse_refine_query의 '
        '자연어 파싱 기능이 핵심 연동 포인트입니다.'
    )

    service_mapping = [
        ('LangChainService', '대화 이력 관리', 'get_conversation_history(), save_conversation()'),
        ('LangChainService', 'Intent Classification', 'Function Calling 스키마 확장'),
        ('parse_refine_query', '검색 조건 파싱', 'target_categories, filters 추출'),
        ('OpenSearchService', '직접 검색', 'hybrid_search() 텍스트 전용 검색'),
        ('RedisService', '세션 상태', 'TTL_CONVERSATION=2시간 활용'),
    ]
    add_table(doc, '', service_mapping, ['서비스', '용도', '메서드/활용'])

    doc.add_page_break()

    # ================================================================
    # 3. 데이터 모델 확장
    # ================================================================
    doc.add_heading('3. 데이터 모델 확장', level=1)

    doc.add_heading('3.1 신규 모델 (선택적)', level=2)

    doc.add_paragraph(
        'Redis 기반 세션 관리를 사용하면 별도 모델이 필요 없습니다. '
        '하지만 대화 이력 분석이 필요한 경우 아래 모델을 추가할 수 있습니다.'
    )

    chat_session_model = '''
# agents/models.py (선택적 - Redis만 사용 시 불필요)

class ChatSession(BaseSoftDeleteModel):
    """에이전트 대화 세션"""
    user = models.ForeignKey(User, on_delete=models.CASCADE)

    # 현재 상태
    current_analysis_id = models.IntegerField(null=True)
    current_intent = models.CharField(max_length=50, null=True)

    # 메타데이터
    created_at = models.DateTimeField(auto_now_add=True)
    last_activity = models.DateTimeField(auto_now=True)

    class Meta:
        db_table = 'chat_session'


class ChatMessage(BaseSoftDeleteModel):
    """대화 메시지 이력"""
    session = models.ForeignKey(ChatSession, on_delete=models.CASCADE)
    role = models.CharField(max_length=20)  # user / assistant
    content = models.TextField()

    # 첨부 이미지
    image_url = models.URLField(null=True)

    # 메타데이터
    intent = models.CharField(max_length=50, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        db_table = 'chat_message'
'''
    add_code_block(doc, chat_session_model, '선택적 모델 (MySQL 저장 시):')

    doc.add_heading('3.2 기존 모델 활용', level=2)

    existing_models = [
        ('UploadedImage', '사용자 업로드 이미지', 'analysis_id 연결'),
        ('ImageAnalysis', '분석 작업 상태', 'status 폴링'),
        ('DetectedObject', '감지된 객체', 'bbox, category'),
        ('ObjectProductMapping', '검색 결과', 'product_id, confidence'),
        ('UserImage', '피팅용 전신 사진', 'user_image_url'),
        ('FittingImage', '피팅 결과', 'fitting_image_url'),
        ('CartItem', '장바구니', '에이전트가 직접 생성'),
        ('Order', '주문', '에이전트가 직접 생성'),
        ('SelectedProduct', '선택 상품', 'size_code 포함'),
    ]
    add_table(doc, '', existing_models, ['모델', '역할', '에이전트 활용'])

    doc.add_page_break()

    # ================================================================
    # 4. API 엔드포인트 설계
    # ================================================================
    doc.add_heading('4. API 엔드포인트 설계', level=1)

    doc.add_heading('4.1 채팅 API', level=2)

    chat_api_spec = '''
POST /api/v1/chat
Authorization: Bearer {access_token}
Content-Type: multipart/form-data (이미지 포함 시) | application/json

Request Body:
{
  "message": "검은색 자켓 찾아줘",     // 필수
  "session_id": "uuid-string",         // 선택, 없으면 새 세션
  "image": <binary>                     // 선택, 참조 이미지
}

Response (200 OK):
{
  "session_id": "uuid-string",
  "response": {
    "text": "검은색 자켓을 찾았어요!\\n\\n1. 나이키 바람막이 - ₩89,000...",
    "type": "search_results",           // text | search_results | fitting | cart | order
    "data": {
      "products": [
        {
          "index": 1,
          "product_id": 123,
          "brand_name": "나이키",
          "product_name": "바람막이 자켓",
          "selling_price": 89000,
          "image_url": "https://...",
          "sizes": ["S", "M", "L"]
        }
      ],
      "total_count": 5
    },
    "suggestions": [                    // 후속 액션 제안
      {"label": "피팅해볼까요?", "action": "fitting"},
      {"label": "다른 색상으로", "action": "refine"}
    ]
  },
  "context": {
    "current_analysis_id": 456,
    "has_search_results": true,
    "has_user_image": true
  }
}

Error Response:
{
  "error": "message_required",
  "detail": "메시지를 입력해주세요."
}
'''
    add_code_block(doc, chat_api_spec, 'POST /api/v1/chat:')

    doc.add_heading('4.2 세션 관리 API (선택)', level=2)

    session_apis = [
        ('GET /api/v1/chat/sessions', '세션 목록', '최근 대화 세션 목록'),
        ('GET /api/v1/chat/sessions/{id}', '세션 상세', '대화 이력 조회'),
        ('DELETE /api/v1/chat/sessions/{id}', '세션 삭제', '대화 초기화'),
    ]
    add_table(doc, '', session_apis, ['엔드포인트', '기능', '설명'])

    doc.add_page_break()

    # ================================================================
    # 5. 에이전트 핵심 로직
    # ================================================================
    doc.add_heading('5. 에이전트 핵심 로직', level=1)

    doc.add_heading('5.1 Intent Classification', level=2)

    doc.add_paragraph(
        'LangChainService의 Function Calling을 확장하여 Intent를 분류합니다. '
        '기존 parse_refine_query 스키마를 확장합니다.'
    )

    intent_schema = '''
INTENT_CLASSIFICATION_SCHEMA = {
    "name": "classify_intent",
    "description": "사용자 메시지의 의도를 분류합니다",
    "parameters": {
        "type": "object",
        "properties": {
            "intent": {
                "type": "string",
                "enum": ["search", "fitting", "commerce", "general", "compound"],
                "description": "주요 의도"
            },
            "sub_intent": {
                "type": "string",
                "enum": [
                    # search
                    "new_search", "refine", "similar", "cross_recommend",
                    # fitting
                    "single_fit", "batch_fit", "compare_fit",
                    # commerce
                    "add_cart", "view_cart", "remove_cart",
                    "size_recommend", "checkout", "order_status",
                    # general
                    "greeting", "help", "feedback", "out_of_scope"
                ]
            },
            "references": {
                "type": "object",
                "properties": {
                    "type": {"type": "string", "enum": ["index", "attribute", "temporal", "none"]},
                    "indices": {"type": "array", "items": {"type": "integer"}},
                    "attribute": {"type": "string"}
                }
            },
            "has_image": {"type": "boolean"},
            "requires_context": {"type": "boolean"},
            "search_params": {
                "type": "object",
                "description": "검색 파라미터 (기존 parse_refine_query 스키마 활용)",
                "properties": {
                    "target_categories": {"type": "array"},
                    "color_filter": {"type": "string"},
                    "brand_filter": {"type": "string"},
                    "style_vibe": {"type": "string"},
                    "price_sort": {"type": "string"}
                }
            }
        },
        "required": ["intent", "sub_intent"]
    }
}
'''
    add_code_block(doc, intent_schema, 'Intent Classification 스키마:')

    doc.add_heading('5.2 서브 에이전트 라우팅', level=2)

    routing_logic = '''
# agents/orchestrator.py

class MainOrchestrator:
    def __init__(self, user_id: int, session_id: str):
        self.user_id = user_id
        self.session_id = session_id
        self.langchain = get_langchain_service()
        self.redis = get_redis_service()

        # 서브 에이전트 초기화
        self.search_agent = SearchAgent(user_id)
        self.fitting_agent = FittingAgent(user_id)
        self.commerce_agent = CommerceAgent(user_id)

    async def process_message(self, message: str, image: bytes = None) -> dict:
        # 1. 세션 컨텍스트 로드
        context = self.load_context()

        # 2. Intent Classification
        intent_result = await self.classify_intent(message, image, context)

        # 3. 라우팅 및 처리
        response = await self.route_to_agent(intent_result, message, image, context)

        # 4. 세션 업데이트
        self.save_context(context, intent_result, response)

        return response

    async def route_to_agent(self, intent_result, message, image, context):
        intent = intent_result['intent']
        sub_intent = intent_result['sub_intent']

        # 복합 의도 처리
        if intent == 'compound':
            return await self.handle_compound(intent_result, message, image, context)

        # 단일 의도 라우팅
        if intent == 'search':
            return await self.search_agent.handle(sub_intent, message, image, context)

        elif intent == 'fitting':
            # 전제조건 확인
            if not context.get('has_search_results'):
                return self.ask_for_search()
            if not context.get('has_user_image'):
                return self.ask_for_user_image()
            return await self.fitting_agent.handle(sub_intent, context)

        elif intent == 'commerce':
            return await self.commerce_agent.handle(sub_intent, message, context)

        else:  # general
            return await self.handle_general(sub_intent, message)
'''
    add_code_block(doc, routing_logic, '라우팅 로직:')

    doc.add_heading('5.3 응답 생성', level=2)

    response_generation = '''
# agents/response_builder.py

class ResponseBuilder:
    """응답 포맷 생성"""

    @staticmethod
    def search_results(products: list, message: str = None) -> dict:
        # 검색 결과 포맷
        text_lines = [message or "찾은 상품이에요:"]
        for i, p in enumerate(products, 1):
            text_lines.append(f"\\n{i}. {p['brand_name']} {p['product_name']} - ₩{p['selling_price']:,}")

        return {
            "text": "\\n".join(text_lines),
            "type": "search_results",
            "data": {
                "products": [
                    {
                        "index": i,
                        "product_id": p['product_id'],
                        "brand_name": p['brand_name'],
                        "product_name": p['product_name'],
                        "selling_price": p['selling_price'],
                        "image_url": p['image_url'],
                        "sizes": p.get('sizes', [])
                    }
                    for i, p in enumerate(products, 1)
                ],
                "total_count": len(products)
            },
            "suggestions": [
                {"label": "피팅해볼까요?", "action": "fitting"},
                {"label": "장바구니에 담기", "action": "add_cart"},
                {"label": "다른 조건으로", "action": "refine"}
            ]
        }

    @staticmethod
    def fitting_result(fitting_url: str, product: dict) -> dict:
        return {
            "text": f"{product['product_name']} 피팅 결과예요!",
            "type": "fitting",
            "data": {
                "fitting_image_url": fitting_url,
                "product": product
            },
            "suggestions": [
                {"label": "장바구니에 담기", "action": "add_cart"},
                {"label": "다른 상품 피팅", "action": "fitting"},
                {"label": "주문하기", "action": "checkout"}
            ]
        }

    @staticmethod
    def error_message(error_type: str, detail: str, suggestions: list = None) -> dict:
        return {
            "text": detail,
            "type": "error",
            "data": {"error_type": error_type},
            "suggestions": suggestions or []
        }
'''
    add_code_block(doc, response_generation, '응답 생성:')

    doc.add_page_break()

    # ================================================================
    # 6. 기존 서비스 연동 상세
    # ================================================================
    doc.add_heading('6. 기존 서비스 연동 상세', level=1)

    doc.add_heading('6.1 Search Agent ↔ 기존 분석 파이프라인', level=2)

    search_agent_code = '''
# agents/sub_agents/search_agent.py

class SearchAgent:
    """검색 에이전트 - 기존 분석 파이프라인 활용"""

    def __init__(self, user_id: int):
        self.user_id = user_id
        self.langchain = get_langchain_service()
        self.opensearch = get_opensearch_client()
        self.redis = get_redis_service()

    async def handle(self, sub_intent: str, message: str, image: bytes, context: dict):
        if sub_intent == 'new_search':
            if image:
                # 이미지 기반 검색 → 기존 분석 파이프라인 활용
                return await self.image_search(image, message, context)
            else:
                # 텍스트 기반 검색 → OpenSearch 직접 검색
                return await self.text_search(message, context)

        elif sub_intent == 'refine':
            # 기존 parse_refine_query 활용 (핵심!)
            return await self.refine_search(message, context)

        elif sub_intent == 'similar':
            return await self.similar_search(context)

        elif sub_intent == 'cross_recommend':
            return await self.cross_category_search(message, context)

    async def image_search(self, image: bytes, message: str, context: dict):
        """이미지 기반 검색 - 기존 API 활용"""
        import base64

        # 1. 이미지 업로드 (기존 API)
        from analyses.views import UploadedImageView
        # 또는 직접 서비스 호출

        # 2. 분석 시작 (기존 Celery Task)
        from analyses.tasks.analysis import process_image_analysis

        # 간소화: 내부적으로 기존 로직 호출
        uploaded_image = await self._upload_image(image)
        analysis = await self._create_analysis(uploaded_image.id)

        # 3. 폴링하여 결과 대기
        results = await self._wait_for_analysis(analysis.id)

        # 4. 컨텍스트 업데이트
        context['current_analysis_id'] = analysis.id
        context['search_results'] = results
        context['has_search_results'] = True

        return ResponseBuilder.search_results(results)

    async def refine_search(self, message: str, context: dict):
        """자연어 재분석 - 기존 parse_refine_query 활용 (핵심!)"""
        analysis_id = context.get('current_analysis_id')
        if not analysis_id:
            return ResponseBuilder.error_message(
                'no_context',
                '먼저 이미지를 업로드하거나 상품을 검색해주세요.'
            )

        # 기존 parse_refine_query_task 활용!
        from analyses.tasks.refine import parse_refine_query_task

        result = parse_refine_query_task.delay(
            query=message,
            available_categories=['top', 'bottom', 'outer', 'shoes', 'bag', 'hat'],
            analysis_id=analysis_id,
            use_v2=True
        ).get()  # 동기 대기 또는 비동기 처리

        # 결과 가져오기
        from analyses.models import ObjectProductMapping
        mappings = ObjectProductMapping.objects.filter(
            detected_object__uploaded_image__imageanalysis=analysis_id
        ).select_related('product')[:5]

        products = [self._mapping_to_product(m) for m in mappings]
        context['search_results'] = products

        return ResponseBuilder.search_results(
            products,
            f"'{message}' 조건으로 다시 찾았어요:"
        )

    async def text_search(self, message: str, context: dict):
        """텍스트 기반 검색 - OpenSearch 직접 활용"""
        # 1. 검색 조건 파싱 (LangChain)
        params = await self._parse_search_params(message)

        # 2. 텍스트 임베딩 생성
        from services import get_embedding_service
        embedding_service = get_embedding_service()
        text_embedding = embedding_service.get_text_embedding(message)

        # 3. OpenSearch 하이브리드 검색
        results = self.opensearch.hybrid_search(
            embedding=text_embedding,
            query=message,
            category=params.get('target_categories', [None])[0],
            filters={
                'color': params.get('color_filter'),
                'brand': params.get('brand_filter'),
                'style': params.get('style_vibe'),
            },
            k=30
        )

        context['search_results'] = results
        context['has_search_results'] = True

        return ResponseBuilder.search_results(results)
'''
    add_code_block(doc, search_agent_code, 'Search Agent 구현:')

    doc.add_heading('6.2 Fitting Agent ↔ 기존 피팅 시스템', level=2)

    fitting_agent_code = '''
# agents/sub_agents/fitting_agent.py

class FittingAgent:
    """피팅 에이전트 - 기존 피팅 API 활용"""

    def __init__(self, user_id: int):
        self.user_id = user_id

    async def handle(self, sub_intent: str, context: dict):
        if sub_intent == 'single_fit':
            return await self.single_fitting(context)
        elif sub_intent == 'batch_fit':
            return await self.batch_fitting(context)
        elif sub_intent == 'compare_fit':
            return await self.compare_fitting(context)

    async def single_fitting(self, context: dict):
        """단일 피팅 - 기존 API 활용"""
        # 1. 선택된 상품 확인
        selected = context.get('selected_product')
        if not selected:
            # 인덱스로 선택 요청
            products = context.get('search_results', [])
            if len(products) == 1:
                selected = products[0]
            else:
                return ResponseBuilder.ask_selection(
                    "어떤 상품을 피팅해볼까요?",
                    products
                )

        # 2. 사용자 이미지 확인
        from fittings.models import UserImage
        user_image = UserImage.objects.filter(
            user_id=self.user_id,
            is_deleted=False
        ).order_by('-created_at').first()

        if not user_image:
            return ResponseBuilder.error_message(
                'no_user_image',
                '피팅을 위해 전신 사진이 필요해요. 사진을 등록해주시겠어요?',
                suggestions=[{"label": "전신 사진 등록", "action": "upload_user_image"}]
            )

        # 3. 피팅 요청 (기존 API 활용)
        from fittings.models import FittingImage
        from fittings.tasks import process_fitting_task

        # 캐시 확인
        existing = FittingImage.objects.filter(
            user_image=user_image,
            product_id=selected['product_id'],
            fitting_image_status='DONE',
            is_deleted=False
        ).first()

        if existing:
            return ResponseBuilder.fitting_result(
                existing.fitting_image_url,
                selected
            )

        # 새 피팅 생성
        fitting = FittingImage.objects.create(
            user_image=user_image,
            product_id=selected['product_id'],
            fitting_image_status='PENDING'
        )

        # Celery 태스크 실행
        process_fitting_task.delay(fitting.id)

        # 폴링 정보 반환
        return {
            "text": "피팅 중이에요... 잠시만 기다려주세요!",
            "type": "fitting_pending",
            "data": {
                "fitting_id": fitting.id,
                "status_url": f"/api/v1/fitting-images/{fitting.id}/status"
            }
        }

    async def batch_fitting(self, context: dict):
        """배치 피팅 - 여러 상품 동시 피팅"""
        products = context.get('search_results', [])

        # 병렬 피팅 요청
        fitting_ids = []
        for product in products[:5]:  # 최대 5개
            context['selected_product'] = product
            result = await self.single_fitting(context)
            if result.get('data', {}).get('fitting_id'):
                fitting_ids.append(result['data']['fitting_id'])

        return {
            "text": f"{len(fitting_ids)}개 상품 피팅 중이에요...",
            "type": "batch_fitting_pending",
            "data": {"fitting_ids": fitting_ids}
        }
'''
    add_code_block(doc, fitting_agent_code, 'Fitting Agent 구현:')

    doc.add_heading('6.3 Commerce Agent ↔ 기존 주문 시스템', level=2)

    commerce_agent_code = '''
# agents/sub_agents/commerce_agent.py

class CommerceAgent:
    """커머스 에이전트 - 기존 Cart/Order 모델 활용"""

    def __init__(self, user_id: int):
        self.user_id = user_id

    async def handle(self, sub_intent: str, message: str, context: dict):
        if sub_intent == 'add_cart':
            return await self.add_to_cart(message, context)
        elif sub_intent == 'view_cart':
            return await self.view_cart()
        elif sub_intent == 'remove_cart':
            return await self.remove_from_cart(message, context)
        elif sub_intent == 'size_recommend':
            return await self.recommend_size(message, context)
        elif sub_intent == 'checkout':
            return await self.checkout(context)
        elif sub_intent == 'order_status':
            return await self.order_status(message)

    async def add_to_cart(self, message: str, context: dict):
        """장바구니 추가 - 기존 모델 활용"""
        from orders.models import CartItem
        from analyses.models import SelectedProduct

        # 1. 선택 상품 확인
        selected = context.get('selected_product')
        if not selected:
            # 인덱스 파싱 시도
            index = self._parse_index(message)
            products = context.get('search_results', [])
            if index and index <= len(products):
                selected = products[index - 1]
            else:
                return ResponseBuilder.ask_selection(
                    "어떤 상품을 담을까요?",
                    products
                )

        # 2. 사이즈 파싱
        size = self._parse_size(message)
        if not size:
            sizes = selected.get('sizes', [])
            if sizes:
                return ResponseBuilder.ask_size(
                    f"{selected['product_name']}의 사이즈를 선택해주세요:",
                    sizes
                )
            # 사이즈 없는 상품
            size = 'FREE'

        # 3. SelectedProduct 생성/조회
        from products.models import SizeCode
        size_code = SizeCode.objects.filter(
            product_id=selected['product_id'],
            size_value=size
        ).first()

        selected_product, _ = SelectedProduct.objects.get_or_create(
            product_id=selected['product_id'],
            size_code=size_code
        )

        # 4. CartItem 생성
        cart_item, created = CartItem.objects.get_or_create(
            user_id=self.user_id,
            selected_product=selected_product,
            is_deleted=False,
            defaults={'quantity': 1}
        )

        if not created:
            cart_item.quantity += 1
            cart_item.save()

        return ResponseBuilder.cart_added(selected, size, cart_item.quantity)

    async def view_cart(self):
        """장바구니 조회 - 기존 API 활용"""
        from orders.models import CartItem

        items = CartItem.objects.filter(
            user_id=self.user_id,
            is_deleted=False
        ).select_related('selected_product__product', 'selected_product__size_code')

        if not items:
            return {
                "text": "장바구니가 비어있어요. 상품을 찾아볼까요?",
                "type": "cart_empty",
                "suggestions": [{"label": "상품 검색", "action": "search"}]
            }

        return ResponseBuilder.cart_list(items)

    async def checkout(self, context: dict):
        """주문 생성 - 기존 Order 모델 활용"""
        from orders.models import CartItem, Order, OrderItem
        from django.db import transaction

        # 1. 장바구니 확인
        cart_items = CartItem.objects.filter(
            user_id=self.user_id,
            is_deleted=False
        ).select_related('selected_product__product')

        if not cart_items:
            return ResponseBuilder.error_message(
                'empty_cart',
                '장바구니가 비어있어요.'
            )

        # 2. 주문 생성
        from users.models import User
        user = User.objects.get(id=self.user_id)

        total_price = sum(
            item.selected_product.product.selling_price * item.quantity
            for item in cart_items
        )

        with transaction.atomic():
            order = Order.objects.create(
                user=user,
                total_price=total_price,
                delivery_address=user.address or '배송지를 입력해주세요'
            )

            for item in cart_items:
                OrderItem.objects.create(
                    order=order,
                    selected_product=item.selected_product,
                    purchased_quantity=item.quantity,
                    price_at_order=item.selected_product.product.selling_price,
                    order_status='PENDING'
                )

            # 장바구니 비우기
            cart_items.update(is_deleted=True)

        return ResponseBuilder.order_created(order, total_price)

    async def recommend_size(self, message: str, context: dict):
        """사이즈 추천"""
        # 사용자 정보 파싱
        height, weight = self._parse_body_info(message)

        selected = context.get('selected_product')
        if not selected:
            return ResponseBuilder.error_message(
                'no_product',
                '어떤 상품의 사이즈를 추천해드릴까요?'
            )

        # 간단한 사이즈 추천 로직
        if height and weight:
            size = self._calculate_size(height, weight, selected.get('category'))
            return {
                "text": f"{height}cm/{weight}kg이시면 {size} 사이즈를 추천드려요!",
                "type": "size_recommendation",
                "data": {
                    "recommended_size": size,
                    "available_sizes": selected.get('sizes', [])
                }
            }

        return {
            "text": "키와 몸무게를 알려주시면 사이즈를 추천해드릴게요!",
            "type": "ask_body_info"
        }
'''
    add_code_block(doc, commerce_agent_code, 'Commerce Agent 구현:')

    doc.add_page_break()

    # ================================================================
    # 7. Redis 상태 관리
    # ================================================================
    doc.add_heading('7. Redis 상태 관리', level=1)

    doc.add_heading('7.1 세션 스키마', level=2)

    redis_schema = '''
# 기존 RedisService 확장

# 키 패턴
agent:session:{session_id}      # 세션 상태 (JSON)
agent:session:{session_id}:turns  # 대화 이력 (List)

# 세션 상태 구조
{
  "session_id": "uuid-string",
  "user_id": 123,

  # 분석 컨텍스트
  "current_analysis_id": 456,
  "current_detected_objects": [
    {"object_id": 1, "category": "top", "bbox": {...}}
  ],

  # 검색 컨텍스트
  "search_results": [
    {"product_id": 789, "brand_name": "Nike", ...}
  ],
  "search_filters": {
    "color": "black",
    "category": "outer"
  },

  # 선택 컨텍스트
  "selected_product": {"product_id": 789, ...},
  "selected_index": 1,

  # 피팅 컨텍스트
  "user_image_id": 101,
  "fitting_results": [
    {"fitting_id": 201, "product_id": 789, "url": "..."}
  ],

  # 장바구니 참조 (실제 데이터는 MySQL)
  "cart_item_count": 3,

  # 메타데이터
  "created_at": "2026-01-20T10:00:00Z",
  "last_activity": "2026-01-20T10:30:00Z"
}

# 대화 이력 (List)
[
  {"role": "user", "content": "검은색 자켓 찾아줘", "timestamp": "..."},
  {"role": "assistant", "content": "검은색 자켓을 찾았어요!", "timestamp": "..."},
  ...
]
'''
    add_code_block(doc, redis_schema, 'Redis 세션 스키마:')

    doc.add_heading('7.2 TTL 정책', level=2)

    ttl_policies = [
        ('agent:session:{id}', '2시간', '기존 TTL_CONVERSATION 활용'),
        ('agent:session:{id}:turns', '2시간', '대화 이력'),
        ('analysis:{id}:status', '30분', '기존 TTL_POLLING'),
        ('fitting:{id}:status', '30분', '기존 TTL_POLLING'),
    ]
    add_table(doc, '', ttl_policies, ['키', 'TTL', '설명'])

    doc.add_page_break()

    # ================================================================
    # 8. Edge Case 처리
    # ================================================================
    doc.add_heading('8. Edge Case 처리', level=1)

    doc.add_heading('8.1 입력 처리', level=2)

    input_edge_cases = [
        ('빈 메시지', '메시지 검증', '"무엇을 도와드릴까요?"'),
        ('이미지만', 'intent=search, has_image=true', '자동 이미지 분석 시작'),
        ('참조 표현', 'references.type 활용', '"이거"→current, "1번"→index'),
        ('오타', 'LLM 자동 교정', '"겁정"→"검정"'),
        ('복합 요청', 'intent=compound', '"찾아서 입어봐"→Search→Fitting'),
    ]
    add_table(doc, '', input_edge_cases, ['케이스', '처리', '비고'])

    doc.add_heading('8.2 대화 흐름', level=2)

    flow_edge_cases = [
        ('컨텍스트 없음', '피팅/장바구니 요청 시', '검색 유도'),
        ('선택 없음', '번호 지정 없이 "담아줘"', '선택 요청'),
        ('세션 만료', 'TTL 2시간 경과', '새 세션 시작 안내'),
        ('중복 요청', '동일 분석 중 재요청', '진행 상태 안내'),
        ('취소', '"취소", "아니야"', '이전 상태 복원'),
    ]
    add_table(doc, '', flow_edge_cases, ['케이스', '상황', '처리'])

    doc.add_heading('8.3 오류 복구', level=2)

    error_recovery = [
        ('API 타임아웃', 'Vision/Claude/피팅', '자동 재시도 3회'),
        ('검색 결과 없음', '조건 불일치', '조건 완화 제안'),
        ('품절', '재고 없음', '대안 상품 추천'),
        ('피팅 실패', 'The New Black 오류', '재시도 제안'),
        ('서버 오류', '500 에러', '일반 오류 안내'),
    ]
    add_table(doc, '', error_recovery, ['케이스', '상황', '처리'])

    doc.add_page_break()

    # ================================================================
    # 9. 구현 단계별 계획
    # ================================================================
    doc.add_heading('9. 구현 단계별 계획', level=1)

    doc.add_heading('9.1 Phase 1: 기반 구축', level=2)

    phase1 = [
        ('1-1', 'agents/ 앱 생성', 'python manage.py startapp agents'),
        ('1-2', 'Intent 스키마 정의', 'agents/schemas.py'),
        ('1-3', 'Orchestrator 기본 구조', 'agents/orchestrator.py'),
        ('1-4', 'Redis 세션 관리', 'services/agent_service.py'),
        ('1-5', 'Chat API 엔드포인트', 'agents/views.py'),
        ('1-6', '기본 테스트', 'agents/tests.py'),
    ]
    add_table(doc, '', phase1, ['#', '작업', '파일'])

    doc.add_heading('9.2 Phase 2: 검색 에이전트', level=2)

    phase2 = [
        ('2-1', 'SearchAgent 구현', 'agents/sub_agents/search_agent.py'),
        ('2-2', '이미지 검색 연동', '기존 process_image_analysis 활용'),
        ('2-3', '텍스트 검색 구현', 'OpenSearchService 직접 활용'),
        ('2-4', 'Refine 연동', '기존 parse_refine_query 활용'),
        ('2-5', '참조 해석 구현', '"이거", "1번" 등'),
        ('2-6', '검색 테스트', '다양한 시나리오 검증'),
    ]
    add_table(doc, '', phase2, ['#', '작업', '설명'])

    doc.add_heading('9.3 Phase 3: 피팅 에이전트', level=2)

    phase3 = [
        ('3-1', 'FittingAgent 구현', 'agents/sub_agents/fitting_agent.py'),
        ('3-2', '전제조건 검사', '검색 결과, 전신 이미지'),
        ('3-3', '단일 피팅 연동', '기존 process_fitting_task 활용'),
        ('3-4', '배치 피팅 구현', '병렬 처리'),
        ('3-5', '피팅 테스트', '다양한 시나리오 검증'),
    ]
    add_table(doc, '', phase3, ['#', '작업', '설명'])

    doc.add_heading('9.4 Phase 4: 커머스 에이전트', level=2)

    phase4 = [
        ('4-1', 'CommerceAgent 구현', 'agents/sub_agents/commerce_agent.py'),
        ('4-2', '장바구니 연동', '기존 CartItem 모델 활용'),
        ('4-3', '사이즈 추천 구현', '신체 정보 파싱'),
        ('4-4', '주문 연동', '기존 Order 모델 활용'),
        ('4-5', '커머스 테스트', '주문 플로우 검증'),
    ]
    add_table(doc, '', phase4, ['#', '작업', '설명'])

    doc.add_heading('9.5 Phase 5: 통합 및 최적화', level=2)

    phase5 = [
        ('5-1', '복합 요청 처리', '"찾아서 입어봐" 등'),
        ('5-2', '오류 복구 강화', '재시도, 대안 제시'),
        ('5-3', '응답 품질 개선', '자연스러운 대화'),
        ('5-4', '성능 최적화', '캐싱, 병렬 처리'),
        ('5-5', '통합 테스트', 'End-to-End 검증'),
        ('5-6', '문서화', 'API 문서, 사용 가이드'),
    ]
    add_table(doc, '', phase5, ['#', '작업', '설명'])

    doc.add_page_break()

    # ================================================================
    # 10. 파일 구조
    # ================================================================
    doc.add_heading('10. 파일 구조', level=1)

    file_structure = '''
Team_G_Backend/
├── agents/                          # 신규 Django App
│   ├── __init__.py
│   ├── apps.py
│   ├── models.py                    # ChatSession, ChatMessage (선택)
│   ├── views.py                     # ChatView
│   ├── serializers.py               # ChatRequest/ResponseSerializer
│   ├── urls.py                      # /api/v1/chat 라우팅
│   │
│   ├── orchestrator.py              # MainOrchestrator
│   ├── schemas.py                   # Intent Classification 스키마
│   ├── response_builder.py          # 응답 포맷 생성
│   │
│   ├── sub_agents/
│   │   ├── __init__.py
│   │   ├── search_agent.py          # SearchAgent
│   │   ├── fitting_agent.py         # FittingAgent
│   │   └── commerce_agent.py        # CommerceAgent
│   │
│   └── tests/
│       ├── __init__.py
│       ├── test_orchestrator.py
│       ├── test_search_agent.py
│       ├── test_fitting_agent.py
│       └── test_commerce_agent.py
│
├── services/
│   └── agent_service.py             # 에이전트 세션 서비스 (신규)
│
├── config/
│   └── urls.py                      # agents.urls 추가
│
└── CLAUDE.md                        # 업데이트 (에이전트 문서화)
'''
    add_code_block(doc, file_structure, '파일 구조:')

    doc.add_heading('설정 파일 수정', level=2)

    config_changes = '''
# config/settings.py
INSTALLED_APPS = [
    ...
    'agents',  # 추가
]

# config/urls.py
urlpatterns = [
    ...
    path('api/v1/', include('agents.urls')),  # 추가
]

# agents/urls.py
from django.urls import path
from .views import ChatView, ChatSessionView

urlpatterns = [
    path('chat', ChatView.as_view(), name='chat'),
    path('chat/sessions', ChatSessionView.as_view(), name='chat_sessions'),
    path('chat/sessions/<str:session_id>', ChatSessionView.as_view(), name='chat_session_detail'),
]
'''
    add_code_block(doc, config_changes, '설정 파일 수정:')

    doc.add_page_break()

    # ================================================================
    # 부록: 전체 흐름도
    # ================================================================
    doc.add_heading('부록: 전체 흐름도', level=1)

    full_flow = '''
┌─────────────────────────────────────────────────────────────────────────────┐
│                        AI 패션 어시스턴트 전체 흐름                           │
├─────────────────────────────────────────────────────────────────────────────┤
│                                                                             │
│  👤 사용자: [셔츠 이미지] + "이거랑 어울리는 바지 찾아줘"                      │
│         │                                                                   │
│         ▼                                                                   │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ POST /api/v1/chat                                                   │   │
│  │ {message: "이거랑 어울리는 바지 찾아줘", image: <binary>}           │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ MainOrchestrator.process_message()                                  │   │
│  │                                                                     │   │
│  │ 1. Redis에서 세션 로드                                               │   │
│  │ 2. Intent Classification (LangChain)                                │   │
│  │    → intent: "search", sub_intent: "cross_recommend"                │   │
│  │ 3. SearchAgent.handle("cross_recommend", ...)                       │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ SearchAgent.cross_category_search()                                 │   │
│  │                                                                     │   │
│  │ 1. 이미지 분석 (기존 파이프라인)                                     │   │
│  │    → VisionService.detect_objects() → 셔츠 감지                     │   │
│  │    → GPT4VService.extract_attributes() → 흰색, 스트라이프            │   │
│  │                                                                     │   │
│  │ 2. 크로스 카테고리 검색                                              │   │
│  │    → OpenSearchService.hybrid_search(                               │   │
│  │         query="어울리는 바지",                                       │   │
│  │         category="bottom",                                          │   │
│  │         context={"셔츠 색상": "흰색", "스타일": "캐주얼"}            │   │
│  │       )                                                             │   │
│  │                                                                     │   │
│  │ 3. Claude 리랭킹                                                    │   │
│  │    → GPT4VService.rerank(candidates, context)                       │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ Response:                                                           │   │
│  │ {                                                                   │   │
│  │   "text": "흰색 스트라이프 셔츠와 어울리는 바지를 찾았어요!\\n\\n     │   │
│  │           1. 네이비 치노팬츠 - ₩59,000\\n                           │   │
│  │           2. 베이지 슬랙스 - ₩45,000\\n                             │   │
│  │           피팅해서 확인해볼까요?",                                   │   │
│  │   "type": "search_results",                                         │   │
│  │   "data": {"products": [...], "total_count": 5},                    │   │
│  │   "suggestions": [{"label": "피팅해볼까요?", "action": "fitting"}]   │   │
│  │ }                                                                   │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  👤 사용자: "1번 입어볼래"                                                  │
│         │                                                                   │
│         ▼                                                                   │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ MainOrchestrator.process_message()                                  │   │
│  │                                                                     │   │
│  │ Intent: "fitting", sub_intent: "single_fit"                         │   │
│  │ references: {type: "index", indices: [1]}                           │   │
│  │                                                                     │   │
│  │ FittingAgent.single_fitting(context)                                │   │
│  │ → context.search_results[0] 선택                                    │   │
│  │ → FittingImage.create(user_image, product)                          │   │
│  │ → process_fitting_task.delay(fitting.id)                            │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  [피팅 완료 후]                                                             │
│                                                                             │
│  👤 사용자: "좋아 담아줘"                                                   │
│         │                                                                   │
│         ▼                                                                   │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │ Intent: "commerce", sub_intent: "add_cart"                          │   │
│  │                                                                     │   │
│  │ CommerceAgent.add_to_cart()                                         │   │
│  │ → SelectedProduct.get_or_create(product, size)                      │   │
│  │ → CartItem.create(user, selected_product)                           │   │
│  └────────────────────────────────┬────────────────────────────────────┘   │
│                                   ▼                                         │
│  Response: "장바구니에 담았어요! 주문하시겠어요?"                            │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, full_flow)

    # ================================================================
    # 저장
    # ================================================================
    output_path = '/Users/ijeong/Desktop/테커/AI_패션_어시스턴트_상세_구현_계획서.docx'
    doc.save(output_path)
    print(f"상세 구현 계획서 저장: {output_path}")
    print(f"파일 크기: {os.path.getsize(output_path):,} bytes")


if __name__ == '__main__':
    create_implementation_plan()
