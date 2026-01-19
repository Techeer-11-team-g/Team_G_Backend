#!/usr/bin/env python3
"""
AI 패션 어시스턴트 에이전트 시스템 - 개발 계획서 생성
현재 프로젝트 기능 기반 상세 구현 계획
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, title, data, headers, level=3):
    if title:
        doc.add_heading(title, level=level)
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    doc.add_paragraph()


def add_code_block(doc, code, title=None):
    if title:
        doc.add_paragraph(title).bold = True
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def create_dev_plan():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)

    # ================================================================
    # 표지
    # ================================================================
    title = doc.add_heading('AI 패션 어시스턴트', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('에이전트 기반 시스템 개발 계획서')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('Team_G Backend 프로젝트 확장')
    doc.add_paragraph('작성일: 2026-01-20')
    doc.add_paragraph('버전: 1.0')
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.add_run('본 문서는 현재 구현된 기능을 기반으로 메인 에이전트 + 서브 에이전트 아키텍처를 적용하기 위한 상세 개발 계획서입니다.')

    doc.add_page_break()

    # ================================================================
    # 목차
    # ================================================================
    doc.add_heading('목차', level=1)
    toc = '''
1. 개요 및 현재 상태
2. 아키텍처 설계
   2.1 전체 아키텍처
   2.2 메인 에이전트 (Orchestrator)
   2.3 서브 에이전트 설계
3. 상세 구현 계획
   3.1 Phase 1: 기반 인프라
   3.2 Phase 2: 서브 에이전트 구현
   3.3 Phase 3: 메인 에이전트 구현
   3.4 Phase 4: 통합 및 최적화
4. Edge Case 처리 계획
5. API 설계
6. 데이터 모델 확장
7. 테스트 계획
8. 배포 전략
'''
    doc.add_paragraph(toc)
    doc.add_page_break()

    # ================================================================
    # 1. 개요 및 현재 상태
    # ================================================================
    doc.add_heading('1. 개요 및 현재 상태', level=1)

    doc.add_heading('1.1 프로젝트 목표', level=2)
    doc.add_paragraph('''
현재 Team_G Backend의 개별 기능들(이미지 분석, 가상 피팅, 주문)을 자연어 대화 인터페이스로 통합하여,
사용자가 하나의 대화 흐름 내에서 검색 → 피팅 → 구매까지 완료할 수 있는 AI 패션 어시스턴트를 구축합니다.
''')

    doc.add_heading('1.2 현재 구현된 기능', level=2)

    add_table(doc, '백엔드 서비스 현황', [
        ('이미지 분석', 'Google Vision → FashionCLIP → OpenSearch', '✅ 구현완료', 'analyses/'),
        ('속성 추출', 'Claude Vision (색상/소재/스타일/패턴)', '✅ 구현완료', 'services/gpt4v_service.py'),
        ('자연어 재분석', 'LangChain Function Calling 파싱', '✅ 구현완료', 'analyses/tasks/refine.py'),
        ('가상 피팅', 'The New Black API (단일 피팅)', '✅ 구현완료', 'fittings/'),
        ('장바구니', 'CartItem CRUD', '✅ 구현완료', 'orders/'),
        ('주문 관리', 'Order/OrderItem 생성/취소', '✅ 구현완료', 'orders/'),
        ('사용자 인증', 'JWT 토큰 기반', '✅ 구현완료', 'users/'),
        ('상태 관리', 'Redis TTL 기반', '✅ 구현완료', 'services/redis_service.py'),
    ], ['기능', '기술', '상태', '위치'])

    add_table(doc, '현재 API 엔드포인트', [
        ('POST', '/api/v1/uploaded-images', '이미지 업로드 (GCS)'),
        ('POST', '/api/v1/analyses', '이미지 분석 시작'),
        ('PATCH', '/api/v1/analyses', '자연어 재분석'),
        ('GET', '/api/v1/analyses/<id>/status', '분석 상태 조회'),
        ('POST', '/api/v1/user-images', '사용자 전신 이미지 업로드'),
        ('POST', '/api/v1/fitting-images', '가상 피팅 요청'),
        ('GET', '/api/v1/fitting-images/<id>/status', '피팅 상태 조회'),
        ('GET/POST', '/api/v1/cart-items', '장바구니 조회/추가'),
        ('POST', '/api/v1/orders', '주문 생성'),
        ('PATCH', '/api/v1/orders/<id>', '주문 취소'),
    ], ['Method', 'Endpoint', '기능'])

    doc.add_heading('1.3 현재 한계점 및 개선 방향', level=2)

    add_table(doc, '', [
        ('이미지 분석 선행 필수', 'Refine은 분석 후에만 가능', '텍스트 전용 검색 지원'),
        ('기능 분리', '검색/피팅/구매가 별도 API', '단일 대화 인터페이스 통합'),
        ('단일 세션', 'analysis_id 범위 내 컨텍스트', '사용자 세션 기반 컨텍스트 확장'),
        ('사이즈 미연동', '수동 사이즈 선택', '신체정보 기반 자동 추천'),
        ('배치 피팅 미지원', '1개씩 순차 요청', '다중 상품 배치 피팅'),
    ], ['현재 한계', '상세', '개선 방향'])

    doc.add_page_break()

    # ================================================================
    # 2. 아키텍처 설계
    # ================================================================
    doc.add_heading('2. 아키텍처 설계', level=1)

    doc.add_heading('2.1 전체 아키텍처', level=2)

    arch_diagram = '''
┌─────────────────────────────────────────────────────────────────────────────┐
│                              클라이언트 (Flutter)                             │
│   ┌────────────────────────────────────────────────────────────────────┐   │
│   │  채팅 UI  │  이미지 업로드  │  상품 카드  │  피팅 결과  │  장바구니  │   │
│   └────────────────────────────────────────────────────────────────────┘   │
└──────────────────────────────────┬──────────────────────────────────────────┘
                                   │ WebSocket / REST API
                                   ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                           API Gateway (Django REST)                          │
│                          POST /api/v1/chat/message                          │
└──────────────────────────────────┬──────────────────────────────────────────┘
                                   │
                                   ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                        🧠 메인 에이전트 (Orchestrator)                         │
│  ┌───────────────────────────────────────────────────────────────────────┐ │
│  │  1. 입력 처리 (텍스트 + 이미지)                                         │ │
│  │  2. 세션 컨텍스트 로드 (Redis)                                          │ │
│  │  3. Intent 분류 (LLM 기반)                                             │ │
│  │  4. 서브 에이전트 라우팅 (단일/순차/병렬)                                 │ │
│  │  5. 결과 통합 및 자연어 응답 생성                                        │ │
│  │  6. 세션 상태 업데이트                                                  │ │
│  └───────────────────────────────────────────────────────────────────────┘ │
│                                   │                                         │
│         ┌─────────────────────────┼─────────────────────────┐               │
│         ▼                         ▼                         ▼               │
│  ┌─────────────────┐     ┌─────────────────┐     ┌─────────────────┐       │
│  │   🔍 Search     │     │   👔 Fitting    │     │   🛒 Commerce   │       │
│  │     Agent       │     │     Agent       │     │     Agent       │       │
│  │                 │     │                 │     │                 │       │
│  │ • 이미지 분석   │     │ • 가상 피팅     │     │ • 사이즈 추천   │       │
│  │ • 유사도 검색   │     │ • 배치 피팅     │     │ • 장바구니 관리 │       │
│  │ • 속성 필터링   │     │ • 컬러 분석     │     │ • 주문 처리     │       │
│  │ • 크로스 추천   │     │ • 비교 모드     │     │ • 가격 안내     │       │
│  └────────┬────────┘     └────────┬────────┘     └────────┬────────┘       │
│           │                       │                       │                 │
└───────────┼───────────────────────┼───────────────────────┼─────────────────┘
            │                       │                       │
            ▼                       ▼                       ▼
┌───────────────────────────────────────────────────────────────────────────────┐
│                              기존 서비스 레이어                                 │
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐          │
│  │ Google      │  │ FashionCLIP │  │ OpenSearch  │  │ Claude      │          │
│  │ Vision API  │  │ Embedding   │  │ k-NN Search │  │ Vision/Text │          │
│  └─────────────┘  └─────────────┘  └─────────────┘  └─────────────┘          │
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐          │
│  │ The New     │  │ Redis       │  │ MySQL       │  │ GCS         │          │
│  │ Black API   │  │ State/Cache │  │ Database    │  │ Storage     │          │
│  └─────────────┘  └─────────────┘  └─────────────┘  └─────────────┘          │
└───────────────────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, arch_diagram, '시스템 아키텍처 다이어그램')

    doc.add_heading('2.2 메인 에이전트 (Orchestrator) 상세', level=2)

    doc.add_heading('2.2.1 Intent 분류 시스템', level=3)

    intent_code = '''
# agents/main_agent.py

class IntentClassifier:
    """사용자 의도 분류기 (Claude 기반)"""

    INTENTS = {
        "search": {
            "description": "상품 검색, 추천, 유사 상품 찾기",
            "keywords": ["찾아", "보여", "검색", "추천", "비슷한", "어울리는", "있어?"],
            "triggers": ["이미지만 전송", "상품명 언급", "카테고리 언급"]
        },
        "fitting": {
            "description": "가상 피팅, 착용 확인, 비교",
            "keywords": ["입어", "피팅", "착용", "어떻게 보여", "비교", "vs"],
            "requires": ["search_results OR product_reference"]
        },
        "commerce": {
            "description": "사이즈 추천, 장바구니, 주문",
            "keywords": ["사이즈", "담아", "장바구니", "주문", "구매", "결제", "배송"],
            "requires": ["product_reference OR fitting_results"]
        },
        "general": {
            "description": "일반 대화, 서비스 안내",
            "keywords": ["안녕", "뭐해", "도움", "사용법"],
            "fallback": True
        },
        "compound": {
            "description": "복합 의도 (검색+피팅, 피팅+구매 등)",
            "patterns": ["~하고 ~해줘", "괜찮으면 담아", "입어보고 주문"]
        }
    }

    async def classify(self, message: str, has_image: bool, context: SessionContext) -> Intent:
        """LLM 기반 의도 분류"""
        prompt = self._build_classification_prompt(message, has_image, context)
        response = await self.llm.generate(prompt)
        return self._parse_intent(response)
'''
    add_code_block(doc, intent_code, 'Intent 분류기 구현')

    doc.add_heading('2.2.2 세션 컨텍스트 관리', level=3)

    session_code = '''
# agents/session.py

@dataclass
class SessionContext:
    """세션 상태 관리 (Redis 저장)"""

    session_id: str
    user_id: int

    # 대화 히스토리 (최근 20턴)
    conversation_history: list[Message]

    # 현재 상태
    current_search_results: list[Product]  # 최근 검색 결과
    selected_products: list[Product]       # 선택된 상품
    fitting_results: list[FittingImage]    # 피팅 결과

    # 사용자 정보
    user_profile: UserProfile              # 키/체중/선호사이즈
    user_preferences: dict                 # 선호 브랜드/스타일

    # 메타데이터
    created_at: datetime
    last_activity: datetime

    TTL_SECONDS = 3600 * 24  # 24시간

    @classmethod
    def from_redis(cls, session_id: str) -> "SessionContext":
        """Redis에서 세션 로드"""
        redis = get_redis_service()
        data = redis.get(f"session:{session_id}")
        return cls(**json.loads(data)) if data else cls.create_new(session_id)

    def save(self):
        """Redis에 세션 저장"""
        redis = get_redis_service()
        redis.setex(f"session:{self.session_id}", self.TTL_SECONDS, json.dumps(asdict(self)))
'''
    add_code_block(doc, session_code, '세션 컨텍스트 구현')

    doc.add_heading('2.2.3 에이전트 라우팅 전략', level=3)

    add_table(doc, '라우팅 규칙', [
        ('이미지만 전송', 'search', '단일', 'Search Agent 호출'),
        ('"찾아줘" + 이미지', 'search', '단일', 'Search Agent 호출'),
        ('"입어볼래"', 'fitting', '단일', 'Fitting Agent (search_results 필요)'),
        ('"담아줘"', 'commerce', '단일', 'Commerce Agent (product 필요)'),
        ('"찾아서 입어봐"', 'compound', '순차', 'Search → Fitting'),
        ('"입어보고 담아줘"', 'compound', '순차', 'Fitting → Commerce'),
        ('"3개 다 입어봐"', 'fitting', '병렬', 'Fitting Agent (배치 모드)'),
        ('"1번 상의, 2번 하의 코디"', 'fitting', '병렬', 'Fitting Agent (조합 모드)'),
    ], ['사용자 입력', 'Intent', '라우팅', '처리'])

    doc.add_page_break()

    doc.add_heading('2.3 서브 에이전트 설계', level=2)

    # Search Agent
    doc.add_heading('2.3.1 Search Agent (검색 에이전트)', level=3)

    search_agent_code = '''
# agents/search_agent.py

class SearchAgent:
    """
    역할: 이미지/텍스트 기반 상품 검색

    기존 활용 서비스:
    - vision_service.py → Google Vision 객체 감지
    - gpt4v_service.py → Claude 속성 추출
    - embedding_service.py → FashionCLIP 임베딩
    - opensearch_client.py → k-NN 검색
    - langchain_service.py → 자연어 파싱
    """

    async def execute(self, request: SearchRequest, context: SessionContext) -> SearchResult:
        """검색 실행"""

        # Case 1: 이미지 기반 검색 (기존 파이프라인 활용)
        if request.image:
            return await self._image_search(request.image, request.query, context)

        # Case 2: 텍스트 전용 검색 (신규)
        elif request.query:
            return await self._text_search(request.query, context)

        # Case 3: 컨텍스트 기반 재검색 (기존 refine 활용)
        elif request.refine_query:
            return await self._refine_search(request.refine_query, context)

    async def _image_search(self, image: bytes, query: str, context: SessionContext):
        """
        기존 process_image_analysis 태스크 활용

        1. Google Vision → 객체 감지 (DetectedObject 생성)
        2. 각 객체별 병렬 처리:
           - 이미지 크롭
           - Claude 속성 추출
           - FashionCLIP 임베딩
           - OpenSearch k-NN 검색
           - Claude 리랭킹
        3. ObjectProductMapping 저장
        """
        # 기존 Celery 태스크 호출
        task = process_image_analysis.delay(
            uploaded_image_id=...,
            image_data=base64.b64encode(image).decode()
        )
        return await self._wait_for_result(task)

    async def _text_search(self, query: str, context: SessionContext):
        """
        텍스트 전용 검색 (신규 구현)

        1. LangChain으로 쿼리 파싱 (카테고리, 색상, 브랜드 등)
        2. 텍스트 임베딩 생성 (FashionCLIP text encoder)
        3. OpenSearch 하이브리드 검색 (벡터 + 필터)
        """
        # 쿼리 파싱
        parsed = await self.langchain_service.parse_search_query(query)

        # 텍스트 임베딩
        embedding = self.embedding_service.get_text_embedding(parsed.description)

        # 하이브리드 검색
        results = self.opensearch.hybrid_search(
            embedding=embedding,
            category=parsed.category,
            filters={
                "color": parsed.color,
                "brand": parsed.brand,
                "price_max": parsed.price_max,
                "style": parsed.style
            }
        )
        return SearchResult(products=results, source="text")

    async def _refine_search(self, query: str, context: SessionContext):
        """
        기존 refine 로직 활용 (analyses/tasks/refine.py)

        1. 대화 히스토리 기반 컨텍스트 파싱
        2. 다중 요청 추출 (여러 카테고리 동시 요청)
        3. 각 카테고리별 재검색
        """
        # 기존 Celery 태스크 호출
        task = parse_refine_query_task.delay(
            analysis_id=context.current_analysis_id,
            query=query,
            detected_categories=context.detected_categories
        )
        return await self._wait_for_result(task)
'''
    add_code_block(doc, search_agent_code, 'Search Agent 구현')

    add_table(doc, 'Search Agent 입출력 명세', [
        ('image', 'bytes | None', '참조 이미지 (있으면 이미지 검색)'),
        ('query', 'str', '자연어 쿼리'),
        ('refine_query', 'str | None', '재검색 쿼리'),
        ('context', 'SessionContext', '세션 컨텍스트'),
        ('---', '---', '---'),
        ('products', 'list[Product]', '검색된 상품 목록 (유사도순)'),
        ('detected_objects', 'list[DetectedObject]', '검출된 객체 (이미지 검색시)'),
        ('match_reasons', 'list[str]', '매칭 이유 설명'),
        ('source', 'str', '검색 소스 (image/text/refine)'),
    ], ['필드', '타입', '설명'])

    doc.add_page_break()

    # Fitting Agent
    doc.add_heading('2.3.2 Fitting Agent (피팅 에이전트)', level=3)

    fitting_agent_code = '''
# agents/fitting_agent.py

class FittingAgent:
    """
    역할: 가상 피팅 및 스타일 시뮬레이션

    기존 활용 서비스:
    - fashn_service.py → The New Black API
    - fittings/tasks.py → process_fitting_task

    신규 기능:
    - 배치 피팅 (다중 상품)
    - 비교 모드 (나란히 비교)
    - 코디 조합 (상의+하의 합성)
    """

    async def execute(self, request: FittingRequest, context: SessionContext) -> FittingResult:
        """피팅 실행"""

        # 사용자 전신 이미지 확인
        user_image = await self._get_user_image(context.user_id)
        if not user_image:
            return FittingResult(
                success=False,
                error="USER_IMAGE_REQUIRED",
                message="가상 피팅을 위해 전신 사진이 필요해요. 먼저 전신 사진을 등록해주세요."
            )

        # Case 1: 단일 피팅
        if len(request.products) == 1:
            return await self._single_fitting(user_image, request.products[0])

        # Case 2: 배치 피팅 (병렬 처리)
        elif request.batch_mode:
            return await self._batch_fitting(user_image, request.products)

        # Case 3: 비교 피팅
        elif request.compare_mode:
            return await self._compare_fitting(user_image, request.products)

        # Case 4: 코디 조합 피팅 (신규)
        elif request.outfit_mode:
            return await self._outfit_fitting(user_image, request.outfit)

    async def _single_fitting(self, user_image: UserImage, product: Product):
        """
        기존 process_fitting_task 활용

        1. 캐시 확인 (동일 조합 재사용)
        2. The New Black API 호출
        3. 결과 저장 및 반환
        """
        # 캐시 확인
        cached = FittingImage.objects.filter(
            user_image=user_image,
            product=product,
            fitting_image_status='DONE'
        ).first()

        if cached:
            return FittingResult(fitting_images=[cached], from_cache=True)

        # 신규 피팅 요청
        fitting = FittingImage.objects.create(
            user_image=user_image,
            product=product,
            fitting_image_status='PENDING'
        )

        # Celery 태스크 호출
        task = process_fitting_task.delay(fitting.id)
        return await self._wait_for_result(task)

    async def _batch_fitting(self, user_image: UserImage, products: list[Product]):
        """
        배치 피팅 (신규) - Celery Group 활용

        여러 상품을 병렬로 피팅 처리
        """
        from celery import group

        tasks = group([
            process_fitting_task.s(
                FittingImage.objects.create(
                    user_image=user_image,
                    product=product,
                    fitting_image_status='PENDING'
                ).id
            )
            for product in products
        ])

        result = tasks.apply_async()
        return await self._wait_for_group_result(result)

    async def _compare_fitting(self, user_image: UserImage, products: list[Product]):
        """
        비교 피팅 - 결과를 나란히 제공
        """
        results = await self._batch_fitting(user_image, products)

        # 비교 분석 추가
        comparison = await self._analyze_comparison(results.fitting_images)

        return FittingResult(
            fitting_images=results.fitting_images,
            comparison=comparison  # 컬러매칭, 스타일매칭 점수
        )
'''
    add_code_block(doc, fitting_agent_code, 'Fitting Agent 구현')

    add_table(doc, 'Fitting Agent 입출력 명세', [
        ('products', 'list[Product]', '피팅할 상품 목록'),
        ('batch_mode', 'bool', '배치 모드 (다중 상품 동시 피팅)'),
        ('compare_mode', 'bool', '비교 모드 (결과 나란히)'),
        ('outfit_mode', 'bool', '코디 모드 (상하의 조합)'),
        ('outfit', 'Outfit | None', '코디 조합 정보'),
        ('---', '---', '---'),
        ('fitting_images', 'list[FittingImage]', '피팅 결과 이미지'),
        ('comparison', 'ComparisonResult | None', '비교 분석 결과'),
        ('from_cache', 'bool', '캐시 사용 여부'),
    ], ['필드', '타입', '설명'])

    doc.add_page_break()

    # Commerce Agent
    doc.add_heading('2.3.3 Commerce Agent (커머스 에이전트)', level=3)

    commerce_agent_code = '''
# agents/commerce_agent.py

class CommerceAgent:
    """
    역할: 사이즈 추천, 장바구니, 주문 관리

    기존 활용 서비스:
    - orders/views.py → CartItem, Order API
    - products/models.py → Product, SizeCode

    신규 기능:
    - 사이즈 자동 추천 (신체정보 기반)
    - 자연어 장바구니 관리
    """

    async def execute(self, request: CommerceRequest, context: SessionContext) -> CommerceResult:
        """커머스 액션 실행"""

        if request.action == "size_recommend":
            return await self._recommend_size(request.product, context.user_profile)

        elif request.action == "cart_add":
            return await self._add_to_cart(request.product, request.size, request.quantity, context)

        elif request.action == "cart_remove":
            return await self._remove_from_cart(request.product_id, context)

        elif request.action == "cart_view":
            return await self._view_cart(context)

        elif request.action == "order_create":
            return await self._create_order(context)

        elif request.action == "order_status":
            return await self._check_order_status(request.order_id, context)

    async def _recommend_size(self, product: Product, profile: UserProfile):
        """
        사이즈 추천 (신규)

        1. 사용자 신체정보 (키, 체중, 선호핏)
        2. 브랜드별 사이즈 가이드 참조
        3. 구매 이력 기반 보정
        """
        # 신체 정보
        height = profile.height  # cm
        weight = profile.weight  # kg
        preferred_fit = profile.preferred_fit  # slim/regular/oversized

        # 브랜드 사이즈 가이드 조회
        size_guide = await self._get_brand_size_guide(product.brand_name)

        # 카테고리별 추천 로직
        if product.category in ['top', 'outer']:
            recommended = self._calc_top_size(height, weight, preferred_fit, size_guide)
        elif product.category in ['pants', 'shorts']:
            recommended = self._calc_pants_size(height, weight, preferred_fit, size_guide)
        elif product.category == 'shoes':
            recommended = self._calc_shoe_size(profile.shoe_size)

        # 구매 이력 보정
        history_adjustment = await self._check_purchase_history(
            context.user_id, product.brand_name
        )

        return SizeRecommendation(
            recommended_size=recommended,
            confidence=0.85,
            alternatives=[...],
            reasoning=f"{height}cm, {weight}kg 기준 {preferred_fit} 핏 추천"
        )

    async def _add_to_cart(self, product: Product, size: str, quantity: int, context: SessionContext):
        """
        장바구니 추가 (기존 API 활용)
        """
        # SizeCode 조회
        size_code = SizeCode.objects.filter(
            product=product,
            size_value=size
        ).first()

        if not size_code:
            return CommerceResult(
                success=False,
                error="SIZE_NOT_AVAILABLE",
                message=f"{size} 사이즈는 품절이에요. 다른 사이즈를 선택해주세요."
            )

        # SelectedProduct 생성
        selected = SelectedProduct.objects.create(
            product=product,
            size_code=size_code
        )

        # CartItem 생성 (기존 로직)
        cart_item = CartItem.objects.create(
            user_id=context.user_id,
            selected_product=selected,
            quantity=quantity
        )

        return CommerceResult(
            success=True,
            cart_item=cart_item,
            message=f"{product.product_name} ({size}) {quantity}개를 장바구니에 담았어요."
        )

    async def _create_order(self, context: SessionContext):
        """
        주문 생성 (기존 API 활용)
        """
        # 장바구니 조회
        cart_items = CartItem.objects.filter(
            user_id=context.user_id,
            is_deleted=False
        ).select_related('selected_product__product', 'selected_product__size_code')

        if not cart_items.exists():
            return CommerceResult(
                success=False,
                error="CART_EMPTY",
                message="장바구니가 비어있어요. 상품을 먼저 담아주세요."
            )

        # 총 금액 계산
        total = sum(
            item.selected_product.product.selling_price * item.quantity
            for item in cart_items
        )

        # Order 생성
        order = Order.objects.create(
            user_id=context.user_id,
            total_price=total,
            delivery_address=context.user_profile.default_address
        )

        # OrderItem 생성
        for item in cart_items:
            OrderItem.objects.create(
                order=order,
                selected_product=item.selected_product,
                purchased_quantity=item.quantity,
                price_at_order=item.selected_product.product.selling_price,
                order_status='PENDING'
            )

        # 장바구니 비우기
        cart_items.update(is_deleted=True)

        return CommerceResult(
            success=True,
            order=order,
            message=f"주문이 생성되었어요. 총 {total:,}원입니다."
        )
'''
    add_code_block(doc, commerce_agent_code, 'Commerce Agent 구현')

    doc.add_page_break()

    # ================================================================
    # 3. 상세 구현 계획
    # ================================================================
    doc.add_heading('3. 상세 구현 계획', level=1)

    doc.add_heading('3.1 Phase 1: 기반 인프라 구축', level=2)

    add_table(doc, '태스크 목록', [
        ('P1-1', '세션 컨텍스트 모듈', 'agents/session.py 생성', '필수'),
        ('P1-2', 'Redis 스키마 확장', 'session:{id} 키 추가, TTL 정책', '필수'),
        ('P1-3', '채팅 API 엔드포인트', 'POST /api/v1/chat/message', '필수'),
        ('P1-4', 'WebSocket 지원 (선택)', 'channels 라이브러리 적용', '선택'),
        ('P1-5', '메시지 모델', 'Message, Conversation 모델 생성', '필수'),
        ('P1-6', 'Celery 큐 확장', 'agent 큐 추가', '필수'),
    ], ['ID', '태스크', '상세', '우선순위'])

    phase1_code = '''
# 파일 구조
agents/
├── __init__.py
├── main_agent.py          # Orchestrator
├── search_agent.py        # Search Agent
├── fitting_agent.py       # Fitting Agent
├── commerce_agent.py      # Commerce Agent
├── session.py             # SessionContext
├── intent.py              # IntentClassifier
├── prompts/               # LLM 프롬프트
│   ├── intent_classification.txt
│   ├── response_generation.txt
│   └── error_handling.txt
└── tests/
    ├── test_main_agent.py
    ├── test_search_agent.py
    └── ...

# Redis 키 스키마
session:{session_id}           → SessionContext JSON (TTL: 24h)
session:{session_id}:history   → 대화 히스토리 (TTL: 24h)
session:{session_id}:search    → 검색 결과 캐시 (TTL: 1h)
session:{session_id}:fitting   → 피팅 결과 캐시 (TTL: 1h)

# Celery 큐
CELERY_TASK_ROUTES = {
    'agents.tasks.*': {'queue': 'agent'},
    'analyses.tasks.*': {'queue': 'analysis'},
    'fittings.tasks.*': {'queue': 'fitting'},
}
'''
    add_code_block(doc, phase1_code, 'Phase 1 구현 상세')

    doc.add_heading('3.2 Phase 2: 서브 에이전트 구현', level=2)

    add_table(doc, '태스크 목록', [
        ('P2-1', 'Search Agent 기본', '기존 분석 파이프라인 래핑', '필수'),
        ('P2-2', 'Search Agent 텍스트', '텍스트 전용 검색 추가', '필수'),
        ('P2-3', 'Fitting Agent 기본', '기존 피팅 태스크 래핑', '필수'),
        ('P2-4', 'Fitting Agent 배치', '다중 상품 병렬 피팅', '필수'),
        ('P2-5', 'Commerce Agent 기본', '기존 장바구니/주문 래핑', '필수'),
        ('P2-6', 'Commerce Agent 사이즈', '사이즈 추천 로직', '필수'),
        ('P2-7', 'UserProfile 모델', '신체정보 저장', '필수'),
    ], ['ID', '태스크', '상세', '우선순위'])

    doc.add_heading('3.3 Phase 3: 메인 에이전트 구현', level=2)

    add_table(doc, '태스크 목록', [
        ('P3-1', 'Intent 분류기', 'Claude 기반 의도 분류', '필수'),
        ('P3-2', '에이전트 라우터', '단일/순차/병렬 라우팅', '필수'),
        ('P3-3', '결과 통합기', '다중 에이전트 결과 병합', '필수'),
        ('P3-4', '응답 생성기', '자연어 응답 생성', '필수'),
        ('P3-5', '참조 해석기', '"이거", "그거", "1번" 등 해석', '필수'),
        ('P3-6', '에러 핸들러', '각 에이전트 에러 통합 처리', '필수'),
    ], ['ID', '태스크', '상세', '우선순위'])

    doc.add_heading('3.4 Phase 4: 통합 및 최적화', level=2)

    add_table(doc, '태스크 목록', [
        ('P4-1', 'E2E 테스트', '전체 시나리오 테스트', '필수'),
        ('P4-2', '성능 최적화', '응답 시간 개선', '필수'),
        ('P4-3', '캐싱 전략', '검색/피팅 결과 캐시', '필수'),
        ('P4-4', '모니터링 확장', '에이전트별 메트릭', '필수'),
        ('P4-5', '에러 복구', '자동 재시도, 폴백', '필수'),
        ('P4-6', '문서화', 'API 문서, 운영 가이드', '필수'),
    ], ['ID', '태스크', '상세', '우선순위'])

    doc.add_page_break()

    # ================================================================
    # 4. Edge Case 처리 계획
    # ================================================================
    doc.add_heading('4. Edge Case 처리 계획', level=1)

    doc.add_paragraph('현재 프로젝트의 구현된 기능 범위에 해당하는 Edge Case만 필터링하여 우선순위별로 정리합니다.')

    doc.add_heading('4.1 [P0] Critical - 반드시 처리', level=2)

    add_table(doc, '입력 처리', [
        ('EC-001', '빈 입력', '텍스트/이미지 모두 없음', '안내 메시지 반환'),
        ('EC-002', '초대형 이미지', '>10MB 이미지 업로드', '리사이징 또는 거부'),
        ('EC-003', '지원 불가 형식', 'GIF/BMP/TIFF 등', '형식 안내'),
        ('EC-004', '손상된 이미지', '읽기 불가 파일', '재업로드 요청'),
        ('EC-005', '패션 아이템 없음', '풍경/음식 등 이미지', '패션 아이템 필요 안내'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, 'API 에러', [
        ('EC-010', 'Google Vision 실패', 'API 타임아웃/에러', '재시도 3회 후 에러 반환'),
        ('EC-011', 'Claude API 실패', 'API 에러/Rate Limit', 'GPT 폴백 또는 기본값'),
        ('EC-012', 'OpenSearch 실패', '검색 서버 오류', '필터 기반 폴백'),
        ('EC-013', '피팅 API 실패', 'The New Black 오류', '재시도 후 에러 안내'),
        ('EC-014', 'Redis 연결 끊김', '세션 로드 실패', '새 세션 생성'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, '비즈니스 에러', [
        ('EC-020', '품절', '재고 없는 상품 담기', '대안 상품 추천'),
        ('EC-021', '사이즈 품절', '특정 사이즈 없음', '다른 사이즈 안내'),
        ('EC-022', '결제 실패', '카드 오류 등', '다른 결제수단 안내'),
        ('EC-023', '주문 취소 불가', '배송 시작 후 취소', '상태 안내'),
        ('EC-024', '전신 이미지 없음', '피팅 요청 시', '이미지 등록 유도'),
    ], ['ID', '시나리오', '상황', '처리'])

    doc.add_heading('4.2 [P1] High - 사용자 경험 필수', level=2)

    add_table(doc, '참조 해석', [
        ('EC-030', '"이거"', '직전 항목 참조', '세션에서 마지막 상품 조회'),
        ('EC-031', '"1번"', '인덱스 참조', '검색 결과에서 해당 인덱스'),
        ('EC-032', '"아까 그거"', '시간 참조', '대화 히스토리에서 조회'),
        ('EC-033', '"그 빨간 거"', '속성 참조', '검색 결과에서 색상 필터'),
        ('EC-034', '"장바구니에 담은 거"', '상태 참조', '장바구니 조회'),
        ('EC-035', '참조 대상 없음', '컨텍스트에 해당 없음', '명확화 질문'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, '멀티턴 대화', [
        ('EC-040', '조건 누적', '"바지" → "검정" → "슬림"', '필터 순차 적용'),
        ('EC-041', '조건 교체', '"검정" → "아니 네이비"', '이전 조건 교체'),
        ('EC-042', '주제 전환', '"바지" → "신발도"', '새 카테고리 추가'),
        ('EC-043', '되돌아가기', '"아까 바지로 돌아가"', '이전 상태 복원'),
        ('EC-044', '부정 누적', '"이거 말고" → "저것도 말고"', '제외 목록 누적'),
        ('EC-045', '세션 만료', '24시간 후 재접속', '새 세션 시작'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, '복합 요청', [
        ('EC-050', '검색+피팅', '"찾아서 입어봐"', 'Search → Fitting 순차'),
        ('EC-051', '피팅+구매', '"입어보고 담아줘"', 'Fitting → Commerce 순차'),
        ('EC-052', '조건부', '"있으면 담아줘"', '검색 → 조건 확인 → 담기'),
        ('EC-053', '비교 요청', '"1번이랑 2번 비교"', '병렬 피팅 → 비교 분석'),
        ('EC-054', '배치 요청', '"다 입어봐"', '배치 피팅 처리'),
    ], ['ID', '시나리오', '상황', '처리'])

    doc.add_heading('4.3 [P2] Medium - 품질 향상', level=2)

    add_table(doc, '이미지 품질', [
        ('EC-060', '저해상도', '100px 미만', '고해상도 요청'),
        ('EC-061', '흐릿함', '초점 불량', '재업로드 권장'),
        ('EC-062', '어두움', '조명 부족', '분석 시도 후 안내'),
        ('EC-063', '필터 적용', '색상 왜곡', '원본 추정'),
        ('EC-064', '워터마크', '로고 삽입', '워터마크 제외 분석'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, '자연어 처리', [
        ('EC-070', '오타', '"겁정", "바즤"', '자동 교정'),
        ('EC-071', '띄어쓰기', '"검정색바지"', '토큰 분리'),
        ('EC-072', '영한 혼용', '"black 바지"', '통합 처리'),
        ('EC-073', '약어', '"나키", "아디"', '브랜드 매핑'),
        ('EC-074', '동의어', '"바지/팬츠"', '동의어 사전'),
        ('EC-075', '초성', '"ㅂㅈ"', '추정 또는 확인'),
    ], ['ID', '시나리오', '상황', '처리'])

    add_table(doc, '사이즈 추천', [
        ('EC-080', '신체정보 없음', '프로필 미입력', '정보 요청'),
        ('EC-081', '브랜드 가이드 없음', '사이즈표 없음', '일반 추천'),
        ('EC-082', '경계 사이즈', 'M/L 사이', '양쪽 안내'),
        ('EC-083', '과거 불만족', '"M 작았어"', '사이즈업 추천'),
    ], ['ID', '시나리오', '상황', '처리'])

    doc.add_heading('4.4 [P3] Low - 추가 개선', level=2)

    add_table(doc, '일반 대화', [
        ('EC-090', '인사', '"안녕"', '인사 응답'),
        ('EC-091', '기능 질문', '"뭘 할 수 있어?"', '기능 안내'),
        ('EC-092', '범위 외', '"날씨 어때?"', '패션 서비스 유도'),
        ('EC-093', '피드백', '"결과가 이상해"', '사과 + 대안'),
        ('EC-094', '칭찬', '"잘 찾네"', '감사 응답'),
    ], ['ID', '시나리오', '상황', '처리'])

    doc.add_page_break()

    # ================================================================
    # 5. API 설계
    # ================================================================
    doc.add_heading('5. API 설계', level=1)

    doc.add_heading('5.1 채팅 API', level=2)

    chat_api = '''
POST /api/v1/chat/message
Content-Type: multipart/form-data

Request:
{
    "message": "string",           // 사용자 메시지 (선택)
    "image": "file",               // 이미지 파일 (선택)
    "session_id": "string"         // 세션 ID (없으면 신규 생성)
}

Response (성공):
{
    "session_id": "abc123",
    "message_id": "msg_001",
    "response": {
        "text": "흰색 스트라이프 셔츠와 어울리는 바지를 찾았어요:",
        "products": [
            {
                "id": 1,
                "name": "네이비 치노팬츠",
                "brand": "ZARA",
                "price": 59000,
                "image_url": "https://...",
                "match_reason": "클래식한 조합"
            }
        ],
        "fitting_images": [],
        "actions": [
            {"type": "fitting", "label": "피팅해보기"},
            {"type": "cart_add", "label": "장바구니 담기"}
        ]
    },
    "intent": "search",
    "processing_time_ms": 2500
}

Response (처리중):
{
    "session_id": "abc123",
    "message_id": "msg_001",
    "status": "processing",
    "progress": 45,
    "estimated_time_seconds": 5
}

Response (에러):
{
    "session_id": "abc123",
    "error": {
        "code": "IMAGE_REQUIRED",
        "message": "가상 피팅을 위해 전신 사진이 필요해요.",
        "action": {
            "type": "upload_user_image",
            "label": "전신 사진 등록하기"
        }
    }
}
'''
    add_code_block(doc, chat_api, '채팅 API 명세')

    doc.add_heading('5.2 세션 API', level=2)

    session_api = '''
GET /api/v1/chat/session/{session_id}
Response:
{
    "session_id": "abc123",
    "user_id": 1,
    "created_at": "2026-01-20T10:00:00Z",
    "last_activity": "2026-01-20T11:30:00Z",
    "conversation_count": 15,
    "current_state": {
        "search_results_count": 9,
        "cart_items_count": 2,
        "fitting_results_count": 3
    }
}

DELETE /api/v1/chat/session/{session_id}
Response:
{
    "message": "세션이 종료되었습니다."
}
'''
    add_code_block(doc, session_api, '세션 API 명세')

    doc.add_heading('5.3 사용자 프로필 API', level=2)

    profile_api = '''
GET /api/v1/users/profile
Response:
{
    "id": 1,
    "email": "user@example.com",
    "profile": {
        "height": 175,
        "weight": 70,
        "shoe_size": 270,
        "preferred_fit": "regular",
        "preferred_brands": ["ZARA", "H&M"],
        "preferred_styles": ["캐주얼", "미니멀"],
        "default_address": "서울시 강남구..."
    }
}

PATCH /api/v1/users/profile
Request:
{
    "height": 175,
    "weight": 70,
    "shoe_size": 270,
    "preferred_fit": "regular"
}
'''
    add_code_block(doc, profile_api, '프로필 API 명세')

    doc.add_page_break()

    # ================================================================
    # 6. 데이터 모델 확장
    # ================================================================
    doc.add_heading('6. 데이터 모델 확장', level=1)

    models_code = '''
# users/models.py - 확장

class UserProfile(models.Model):
    """사용자 신체정보 및 선호도"""
    user = models.OneToOneField(User, on_delete=models.CASCADE, related_name='profile')

    # 신체 정보
    height = models.IntegerField(null=True, help_text='키 (cm)')
    weight = models.IntegerField(null=True, help_text='체중 (kg)')
    shoe_size = models.IntegerField(null=True, help_text='신발 사이즈 (mm)')

    # 선호도
    preferred_fit = models.CharField(
        max_length=20,
        choices=[('slim', 'Slim'), ('regular', 'Regular'), ('oversized', 'Oversized')],
        default='regular'
    )
    preferred_brands = models.JSONField(default=list, help_text='선호 브랜드 목록')
    preferred_styles = models.JSONField(default=list, help_text='선호 스타일 목록')

    # 배송지
    default_address = models.TextField(null=True, blank=True)

    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)


# agents/models.py - 신규

class Conversation(BaseSoftDeleteModel):
    """대화 세션"""
    session_id = models.CharField(max_length=64, unique=True, db_index=True)
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='conversations')

    created_at = models.DateTimeField(auto_now_add=True)
    last_activity = models.DateTimeField(auto_now=True)

    class Meta:
        ordering = ['-last_activity']


class Message(BaseSoftDeleteModel):
    """개별 메시지"""
    conversation = models.ForeignKey(Conversation, on_delete=models.CASCADE, related_name='messages')

    role = models.CharField(max_length=20, choices=[('user', 'User'), ('assistant', 'Assistant')])
    content = models.TextField()

    # 메타데이터
    intent = models.CharField(max_length=50, null=True)
    has_image = models.BooleanField(default=False)
    image_url = models.URLField(null=True, blank=True)

    # 관련 데이터
    related_products = models.JSONField(default=list, help_text='관련 상품 ID 목록')
    related_fittings = models.JSONField(default=list, help_text='관련 피팅 ID 목록')

    processing_time_ms = models.IntegerField(null=True)

    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        ordering = ['created_at']


class BrandSizeGuide(models.Model):
    """브랜드별 사이즈 가이드"""
    brand_name = models.CharField(max_length=100, db_index=True)
    category = models.CharField(max_length=50)  # top, pants, shoes

    # 사이즈 매핑 (JSON)
    # {"S": {"chest": [86, 91], "height": [160, 168]}, ...}
    size_chart = models.JSONField()

    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        unique_together = ['brand_name', 'category']
'''
    add_code_block(doc, models_code, '모델 확장')

    doc.add_page_break()

    # ================================================================
    # 7. 테스트 계획
    # ================================================================
    doc.add_heading('7. 테스트 계획', level=1)

    doc.add_heading('7.1 단위 테스트', level=2)

    add_table(doc, '', [
        ('Intent 분류', 'test_intent_classifier.py', '각 Intent 유형별 정확도'),
        ('Search Agent', 'test_search_agent.py', '이미지/텍스트/재검색'),
        ('Fitting Agent', 'test_fitting_agent.py', '단일/배치/비교'),
        ('Commerce Agent', 'test_commerce_agent.py', '사이즈/장바구니/주문'),
        ('세션 관리', 'test_session.py', 'Redis 저장/로드/만료'),
        ('참조 해석', 'test_reference.py', '"이거", "1번" 등'),
    ], ['테스트 대상', '파일', '범위'])

    doc.add_heading('7.2 통합 테스트', level=2)

    add_table(doc, '', [
        ('E2E-001', '이미지 검색 → 피팅 → 구매', '전체 플로우'),
        ('E2E-002', '텍스트 검색 → 조건 추가 → 피팅', '멀티턴 대화'),
        ('E2E-003', '배치 피팅 → 비교 → 선택', '다중 상품'),
        ('E2E-004', '사이즈 추천 → 장바구니 → 주문', '커머스 플로우'),
        ('E2E-005', '세션 유지 → 만료 → 복구', '세션 관리'),
    ], ['ID', '시나리오', '검증 포인트'])

    doc.add_heading('7.3 성능 테스트', level=2)

    add_table(doc, '', [
        ('응답 시간', '단일 검색 응답', '< 3초', '95th percentile'),
        ('응답 시간', '피팅 요청 → 완료', '< 30초', '단일 피팅'),
        ('응답 시간', '배치 피팅 (5개)', '< 60초', '병렬 처리'),
        ('처리량', '동시 요청', '100 req/s', '검색 API'),
        ('메모리', 'Redis 세션', '< 100MB', '1000 세션'),
    ], ['카테고리', '항목', '목표', '조건'])

    doc.add_page_break()

    # ================================================================
    # 8. 배포 전략
    # ================================================================
    doc.add_heading('8. 배포 전략', level=1)

    doc.add_heading('8.1 단계별 배포', level=2)

    add_table(doc, '', [
        ('1단계', '내부 테스트', 'dev 환경, QA팀', '기능 검증'),
        ('2단계', '베타 출시', 'staging 환경, 제한된 사용자', '사용성 피드백'),
        ('3단계', '점진적 출시', 'prod 환경, 10% → 50% → 100%', '안정성 확보'),
    ], ['단계', '범위', '대상', '목적'])

    doc.add_heading('8.2 롤백 계획', level=2)

    doc.add_paragraph('''
- 기존 API 유지: /api/v1/analyses, /api/v1/fitting-images 등 기존 엔드포인트는 그대로 유지
- Feature Flag: 채팅 API는 Feature Flag로 on/off 제어
- 장애 발생 시: 채팅 기능만 비활성화, 기존 기능은 정상 운영
- 데이터 호환: 기존 모델 변경 없음, 신규 모델만 추가
''')

    doc.add_heading('8.3 모니터링 확장', level=2)

    monitor_code = '''
# 신규 Prometheus 메트릭

# 에이전트별 메트릭
teamg_agent_requests_total{agent="search|fitting|commerce", status="success|error"}
teamg_agent_duration_seconds{agent="search|fitting|commerce"}

# Intent 분류 메트릭
teamg_intent_classification_total{intent="search|fitting|commerce|general|compound"}
teamg_intent_confidence_histogram{intent}

# 세션 메트릭
teamg_active_sessions_gauge
teamg_session_duration_seconds
teamg_messages_per_session_histogram

# 에러 메트릭
teamg_agent_errors_total{agent, error_code}
'''
    add_code_block(doc, monitor_code, '모니터링 메트릭')

    doc.add_page_break()

    # ================================================================
    # 부록: 구현 우선순위 요약
    # ================================================================
    doc.add_heading('부록: 구현 우선순위 요약', level=1)

    summary_table = [
        ('1', 'Phase 1', '기반 인프라', '세션, Redis, 채팅 API', '필수'),
        ('2', 'Phase 2', 'Search Agent', '기존 분석 래핑 + 텍스트 검색', '필수'),
        ('3', 'Phase 2', 'Fitting Agent', '기존 피팅 래핑 + 배치', '필수'),
        ('4', 'Phase 2', 'Commerce Agent', '장바구니/주문 래핑 + 사이즈', '필수'),
        ('5', 'Phase 3', '메인 에이전트', 'Intent 분류 + 라우팅', '필수'),
        ('6', 'Phase 3', '참조 해석', '"이거", "1번" 등', '필수'),
        ('7', 'Phase 4', '성능 최적화', '캐싱, 응답시간', '필수'),
        ('8', 'Phase 4', '모니터링', '에이전트 메트릭', '필수'),
        ('---', '---', '---', '---', '---'),
        ('9', '추가', '비교 피팅', '나란히 비교, 점수', '선택'),
        ('10', '추가', '코디 조합', '상의+하의 합성', '선택'),
        ('11', '추가', 'WebSocket', '실시간 상태', '선택'),
    ]

    add_table(doc, '구현 우선순위', summary_table, ['순번', 'Phase', '항목', '상세', '우선순위'])

    # 저장
    output_path = '/Users/ijeong/Desktop/테커/AI_패션_어시스턴트_에이전트_개발_계획서.docx'
    doc.save(output_path)
    print(f"개발 계획서 저장: {output_path}")
    print(f"파일 크기: {os.path.getsize(output_path):,} bytes")


if __name__ == '__main__':
    create_dev_plan()
