#!/usr/bin/env python3
"""
AI 패션 어시스턴트 에이전트 시스템 - 개발 계획서 (UX 중심)
사용자 경험을 최우선으로 한 정제된 계획서
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def add_table(doc, title, data, headers, level=3):
    if title:
        doc.add_heading(title, level=level)
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    doc.add_paragraph()


def add_code_block(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph()


def create_ux_focused_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)

    # ================================================================
    # 표지
    # ================================================================
    title = doc.add_heading('AI 패션 어시스턴트', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('메인 에이전트 개발 계획서')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('사용자 경험 중심 설계\n').bold = True
    info.add_run('Team_G Backend 프로젝트\n')
    info.add_run('작성일: 2026-01-20')

    doc.add_page_break()

    # ================================================================
    # 목차
    # ================================================================
    doc.add_heading('목차', level=1)
    toc = '''
1. 개요
   1.1 프로젝트 목표
   1.2 핵심 설계 원칙
   1.3 아키텍처 요약

2. 메인 에이전트 (Orchestrator)
   2.1 역할과 책임
   2.2 의도 분류 (Intent Classification)
   2.3 서브 에이전트 라우팅
   2.4 세션 컨텍스트 관리

3. 사용자 여정별 시나리오
   3.1 검색 여정 (Search Journey)
   3.2 피팅 여정 (Fitting Journey)
   3.3 구매 여정 (Commerce Journey)
   3.4 복합 여정 (Combined Journey)

4. 자연어 처리 상세
   4.1 참조 표현 해석
   4.2 멀티턴 대화 처리
   4.3 모호한 입력 처리
   4.4 오류 입력 복구

5. 응답 생성 가이드라인
   5.1 응답 구조
   5.2 톤앤매너
   5.3 액션 버튼 제안
   5.4 오류 메시지

6. Edge Case 처리 우선순위
   6.1 P0: 필수 (서비스 중단 방지)
   6.2 P1: 핵심 UX (자연스러운 대화)
   6.3 P2: 품질 향상
   6.4 P3: 추가 개선

7. 구현 로드맵
'''
    doc.add_paragraph(toc)
    doc.add_page_break()

    # ================================================================
    # 1. 개요
    # ================================================================
    doc.add_heading('1. 개요', level=1)

    doc.add_heading('1.1 프로젝트 목표', level=2)
    doc.add_paragraph(
        '"자연어와 이미지로 대화하며 쇼핑하는 AI 어시스턴트"\n\n'
        '사용자가 복잡한 UI 없이 대화만으로 상품 검색, 가상 피팅, 구매까지 '
        '완료할 수 있는 통합 쇼핑 경험을 제공합니다.'
    )

    doc.add_heading('1.2 핵심 설계 원칙', level=2)

    principles = [
        ('🎯 의도 우선', '사용자가 "무엇을 원하는지"를 정확히 파악하고, 최적의 에이전트로 라우팅'),
        ('💬 자연스러운 대화', '"이거", "아까 그거", "1번" 같은 일상 표현을 자연스럽게 이해'),
        ('🔄 컨텍스트 유지', '이전 대화를 기억하고 연속적인 쇼핑 경험 제공'),
        ('❌ 실패 없는 경험', '어떤 입력에도 막다른 길 없이 대안 제시'),
        ('⚡ 빠른 응답', '핵심 정보 먼저, 상세 정보는 요청 시 제공'),
    ]
    add_table(doc, '', principles, ['원칙', '설명'])

    doc.add_heading('1.3 아키텍처 요약', level=2)

    arch_diagram = '''
┌─────────────────────────────────────────────────────────────────────┐
│                    🧠 메인 에이전트 (Orchestrator)                    │
│                                                                     │
│   [사용자 입력] → Intent 분석 → 서브 에이전트 라우팅 → 응답 생성      │
│                                                                     │
├─────────────────────────────────────────────────────────────────────┤
│         │                    │                    │                 │
│         ▼                    ▼                    ▼                 │
│   ┌───────────┐        ┌───────────┐        ┌───────────┐          │
│   │ 🔍 Search │        │ 👔 Fitting│        │ 🛒 Commerce│          │
│   │   Agent   │        │   Agent   │        │   Agent   │          │
│   └─────┬─────┘        └─────┬─────┘        └─────┬─────┘          │
│         │                    │                    │                 │
│   ┌─────┴─────┐        ┌─────┴─────┐        ┌─────┴─────┐          │
│   │• 이미지분석│        │• 가상피팅 │        │• 사이즈추천│          │
│   │• 유사검색 │        │• 스타일비교│        │• 장바구니 │          │
│   │• 필터링   │        │• 코디피팅 │        │• 주문     │          │
│   └───────────┘        └───────────┘        └───────────┘          │
│                                                                     │
│   [현재 구현된 기술 스택]                                            │
│   • Google Vision API (객체 감지)                                   │
│   • Claude Vision (속성 추출 + 리랭킹)                               │
│   • FashionCLIP (512차원 임베딩)                                    │
│   • OpenSearch k-NN (벡터 검색)                                     │
│   • The New Black API (가상 피팅)                                   │
│   • Redis (세션/상태 관리)                                          │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, arch_diagram)

    doc.add_page_break()

    # ================================================================
    # 2. 메인 에이전트 (Orchestrator)
    # ================================================================
    doc.add_heading('2. 메인 에이전트 (Orchestrator)', level=1)

    doc.add_heading('2.1 역할과 책임', level=2)

    responsibilities = [
        ('입력 처리', '텍스트/이미지 파싱, 세션 컨텍스트 로드'),
        ('의도 분류', '사용자 발화의 의도(Intent) 파악'),
        ('라우팅', '적절한 서브 에이전트로 요청 전달'),
        ('결과 통합', '서브 에이전트 결과를 자연어 응답으로 변환'),
        ('상태 관리', '대화 이력, 검색 결과, 선택 상품 등 세션 유지'),
    ]
    add_table(doc, '', responsibilities, ['책임', '설명'])

    doc.add_heading('2.2 의도 분류 (Intent Classification)', level=2)

    doc.add_paragraph(
        '사용자 입력을 분석하여 적절한 처리 경로를 결정합니다. '
        'Claude LLM의 Function Calling을 활용하여 의도를 분류합니다.'
    )

    intent_schema = '''
{
  "intent": "search | fitting | commerce | general | compound",
  "sub_intent": "string",  // 세부 의도
  "has_image": boolean,
  "references": {
    "type": "index | attribute | temporal | none",
    "value": "any"
  },
  "parameters": {
    // 의도별 파라미터
  }
}
'''
    add_code_block(doc, intent_schema, 'Intent 스키마:')

    # Intent 분류 테이블
    intent_table = [
        ('search', '찾아줘, 보여줘, 검색, 추천, 비슷한, 있어?', '이미지만 전송해도 search'),
        ('fitting', '입어볼래, 피팅, 착용, 비교해줘, 어떻게 보여?', '검색 결과 필요'),
        ('commerce', '담아줘, 장바구니, 사이즈, 주문, 구매', '상품 선택 필요'),
        ('general', '안녕, 도와줘, 뭐해?, 뭘 할 수 있어?', 'LLM 직접 응답'),
        ('compound', '찾아서 입어봐, 입어보고 담아줘', '순차 처리'),
    ]
    add_table(doc, 'Intent 분류 기준', intent_table, ['Intent', '트리거 키워드', '비고'])

    doc.add_heading('2.3 서브 에이전트 라우팅', level=2)

    routing_rules = [
        ('이미지만 전송', 'Search Agent', '자동으로 이미지 분석 시작'),
        ('"찾아줘" + 조건', 'Search Agent', '조건 파싱 후 검색'),
        ('"입어볼래"', 'Fitting Agent', 'current_results 필요, 없으면 안내'),
        ('"담아줘"', 'Commerce Agent', 'selected_product 필요'),
        ('"사이즈?"', 'Commerce Agent', '사용자 정보 확인'),
        ('"찾아서 입어봐"', 'Search → Fitting', '순차 체인'),
        ('"입어보고 담아"', 'Fitting → Commerce', '순차 체인'),
    ]
    add_table(doc, '라우팅 규칙', routing_rules, ['입력 패턴', '라우팅', '비고'])

    doc.add_heading('2.4 세션 컨텍스트 관리', level=2)

    context_schema = '''
SessionContext:
  session_id: str
  user_id: int

  # 대화 상태
  conversation_history: list[Message]  # 최근 20턴
  current_intent: Intent

  # 검색 상태
  current_search_results: list[Product]  # 최근 검색 결과
  search_filters: dict                   # 적용된 필터

  # 선택 상태
  selected_products: list[Product]       # 선택된 상품
  fitting_results: list[FittingImage]    # 피팅 결과

  # 사용자 정보
  user_profile: {height, weight, preferred_fit}

  # TTL
  created_at: datetime
  last_activity: datetime
  ttl: 24시간
'''
    add_code_block(doc, context_schema, '세션 컨텍스트 구조:')

    doc.add_page_break()

    # ================================================================
    # 3. 사용자 여정별 시나리오
    # ================================================================
    doc.add_heading('3. 사용자 여정별 시나리오', level=1)

    doc.add_heading('3.1 검색 여정 (Search Journey)', level=2)

    # 검색 시작 시나리오
    doc.add_heading('3.1.1 검색 시작 방법', level=3)

    search_start = [
        ('이미지만', '[사진 첨부]', '→ 자동으로 이미지 내 아이템 분석 시작'),
        ('이미지+텍스트', '[사진] + "이거 찾아줘"', '→ 이미지 분석 + 의도 확인'),
        ('텍스트만', '"검은색 가죽 자켓 찾아줘"', '→ 조건 파싱 후 텍스트 검색'),
        ('크로스 추천', '[셔츠 사진] + "어울리는 바지"', '→ 이미지 분석 후 다른 카테고리 검색'),
    ]
    add_table(doc, '', search_start, ['유형', '입력 예시', '처리'])

    # 검색 수정 시나리오
    doc.add_heading('3.1.2 검색 결과 수정', level=3)

    search_modify = [
        ('필터 추가', '"검정으로", "나이키로", "면 소재"', '기존 결과에 필터 적용'),
        ('필터 교체', '"아니 네이비로", "검정 말고 흰색"', '이전 필터 교체'),
        ('필터 제거', '"색상 빼고", "조건 없이"', '특정/전체 필터 제거'),
        ('결과 좁히기', '"이 중에서 무지만"', '현재 결과 내 필터'),
        ('결과 넓히기', '"비슷한 것도 더"', '조건 완화 재검색'),
        ('정렬 변경', '"싼 순으로", "비싼 순으로"', 'price_sort 적용'),
        ('더보기', '"더 보여줘", "다음"', '페이지네이션'),
    ]
    add_table(doc, '', search_modify, ['동작', '입력 예시', '처리'])

    # 검색 예외 상황
    doc.add_heading('3.1.3 검색 예외 처리', level=3)

    search_exceptions = [
        ('결과 없음', '"X에 맞는 상품이 없어요. Y 조건으로 찾아볼까요?"', '조건 완화 제안'),
        ('모호한 카테고리', '"어떤 종류의 옷을 찾으시나요? 상의/하의/아우터..."', '카테고리 질문'),
        ('미지원 카테고리', '"아쉽지만 시계는 아직 지원하지 않아요. 상의/하의/신발/가방은 가능해요."', '대안 안내'),
        ('이미지 품질 불량', '"이미지가 흐릿해서 분석이 어려워요. 선명한 사진으로 다시 올려주시겠어요?"', '재업로드 요청'),
        ('패션 아이템 없음', '"이미지에서 패션 아이템을 찾지 못했어요. 옷이 보이는 사진을 올려주세요."', '안내'),
    ]
    add_table(doc, '', search_exceptions, ['상황', '응답 예시', '전략'])

    doc.add_heading('3.2 피팅 여정 (Fitting Journey)', level=2)

    # 피팅 요청 유형
    doc.add_heading('3.2.1 피팅 요청 유형', level=3)

    fitting_types = [
        ('단일 피팅', '"이거 입어볼래", "1번 피팅해줘"', '선택 상품 1개 피팅'),
        ('전체 피팅', '"다 입어봐", "전부 피팅"', '검색 결과 전체 배치 피팅'),
        ('선택 피팅', '"1번이랑 3번만"', '지정 상품만 피팅'),
        ('비교 피팅', '"둘 다 입어서 비교", "A vs B"', '그리드 레이아웃 비교'),
        ('변형 피팅', '"이거 검정으로 피팅"', '색상 변형 검색 후 피팅'),
    ]
    add_table(doc, '', fitting_types, ['유형', '입력 예시', '처리'])

    # 피팅 전제조건
    doc.add_heading('3.2.2 피팅 전제조건 확인', level=3)

    fitting_prereqs = [
        ('전신 이미지 없음', '"피팅을 위해 전신 사진이 필요해요. 사진을 등록해주시겠어요?"', '업로드 유도'),
        ('검색 결과 없음', '"먼저 피팅할 상품을 찾아볼까요? 어떤 옷을 찾으시나요?"', '검색 유도'),
        ('미지원 카테고리', '"아쉽지만 액세서리는 피팅이 어려워요. 상의/하의/신발/가방 피팅은 가능해요!"', '대안 안내'),
        ('피팅 진행 중', '"이전 피팅이 진행 중이에요. 완료되면 바로 알려드릴게요!"', '대기 안내'),
    ]
    add_table(doc, '', fitting_prereqs, ['상황', '응답 예시', '전략'])

    doc.add_heading('3.3 구매 여정 (Commerce Journey)', level=2)

    # 장바구니 관리
    doc.add_heading('3.3.1 장바구니 관리', level=3)

    cart_actions = [
        ('담기', '"담아줘", "1번 담아"', 'CartItem 생성'),
        ('옵션 지정', '"검정 L 2개 담아"', '색상+사이즈+수량 파싱'),
        ('조회', '"장바구니 보여줘"', '목록 + 합계 표시'),
        ('삭제', '"바지 빼줘", "1번 삭제"', 'soft delete'),
        ('수량 변경', '"3개로 바꿔줘"', 'quantity 업데이트'),
        ('비우기', '"전부 비워줘"', '전체 삭제'),
    ]
    add_table(doc, '', cart_actions, ['동작', '입력 예시', '처리'])

    # 사이즈 추천
    doc.add_heading('3.3.2 사이즈 추천', level=3)

    size_scenarios = [
        ('정보 있음', '"175cm 70kg인데 사이즈?"', '즉시 추천 (L 추천, 신뢰도 87%)'),
        ('정보 없음', '"사이즈 뭐가 좋아?"', '신체 정보 요청'),
        ('선호 기반', '"넉넉하게 입고 싶어"', '사이즈업 추천'),
        ('브랜드별', '"이 브랜드는 작게 나와?"', '브랜드 사이즈 가이드 참조'),
        ('재고 확인', '"L 있어?"', 'SizeCode 조회'),
    ]
    add_table(doc, '', size_scenarios, ['상황', '입력 예시', '처리'])

    # 주문
    doc.add_heading('3.3.3 주문 프로세스', level=3)

    order_scenarios = [
        ('장바구니 주문', '"주문할게"', '장바구니 → 주문 확인 → 결제'),
        ('즉시 구매', '"이거 바로 살게"', '선택 상품 → 주문 확인'),
        ('주문 조회', '"주문 내역 보여줘"', '최근 주문 목록'),
        ('배송 조회', '"배송 어디야?"', '배송 상태 조회'),
        ('취소', '"주문 취소해줘"', '취소 가능 여부 확인 → 처리'),
    ]
    add_table(doc, '', order_scenarios, ['상황', '입력 예시', '처리'])

    doc.add_heading('3.4 복합 여정 (Combined Journey)', level=2)

    # 순차 처리
    compound_scenarios = [
        ('찾아서 입어봐', '"이거 비슷한 거 찾아서 입어봐"', 'Search → Fitting'),
        ('입어보고 담아', '"1번 입어보고 괜찮으면 담아줘"', 'Fitting → (확인) → Commerce'),
        ('찾아서 담아', '"검정 바지 찾아서 담아줘"', 'Search → (선택 요청) → Commerce'),
        ('전체 플로우', '"이 코디 찾아서 다 입어보고 괜찮은 거 담아줘"', 'Search → Fitting → Commerce'),
    ]
    add_table(doc, '', compound_scenarios, ['패턴', '입력 예시', '처리 체인'])

    doc.add_page_break()

    # ================================================================
    # 4. 자연어 처리 상세
    # ================================================================
    doc.add_heading('4. 자연어 처리 상세', level=1)

    doc.add_heading('4.1 참조 표현 해석', level=2)

    doc.add_paragraph(
        '사용자가 "이거", "아까 그거", "1번" 같은 참조 표현을 사용할 때 '
        '세션 컨텍스트를 활용하여 정확한 대상을 파악합니다.'
    )

    # 직접 참조
    direct_refs = [
        ('"이거"', 'current_item', '현재 표시 중인 항목'),
        ('"그거"', 'last_mentioned_item', '직전에 언급한 항목'),
        ('"1번", "첫 번째"', 'search_results[0]', '인덱스 참조'),
        ('"마지막 거"', 'search_results[-1]', '마지막 항목'),
        ('"빨간 거"', 'filter(color=red)', '속성으로 필터'),
        ('"나이키 거"', 'filter(brand=nike)', '브랜드로 필터'),
    ]
    add_table(doc, '직접 참조', direct_refs, ['표현', '해석', '설명'])

    # 시간 참조
    temporal_refs = [
        ('"아까"', 'conversation_history[-N]', '이전 턴의 맥락'),
        ('"방금"', 'last_turn', '직전 턴'),
        ('"처음에"', 'first_search_results', '세션 첫 검색'),
        ('"어제", "지난번"', '세션 범위 외', '안내: "현재 대화에서는..."'),
    ]
    add_table(doc, '시간 참조', temporal_refs, ['표현', '해석', '설명'])

    # 상태 참조
    state_refs = [
        ('"검색한 거"', 'search_results', '검색 결과'),
        ('"피팅한 거"', 'fitting_results', '피팅 결과'),
        ('"담은 거"', 'cart_items', '장바구니'),
        ('"주문한 거"', 'orders', '주문 이력'),
    ]
    add_table(doc, '상태 참조', state_refs, ['표현', '해석', '설명'])

    doc.add_heading('4.2 멀티턴 대화 처리', level=2)

    # 조건 누적
    multiturn_patterns = [
        ('조건 누적', '"바지" → "검정" → "슬림"', '순차적으로 필터 추가'),
        ('조건 교체', '"검정" → "아니 네이비로"', '"아니"로 이전 조건 교체 감지'),
        ('주제 전환', '"바지" → "신발도 찾아줘"', '새 카테고리 추가, 기존 유지'),
        ('되돌아가기', '"아까 바지로 돌아가"', '이전 검색 상태 복원'),
        ('결과 좁히기', '"이 중에서 무지만"', '현재 결과 내 필터'),
        ('수정', '"아니 그게 아니라"', '이전 해석 취소'),
    ]
    add_table(doc, '대화 패턴 처리', multiturn_patterns, ['패턴', '예시', '처리'])

    doc.add_heading('4.3 모호한 입력 처리', level=2)

    ambiguous_inputs = [
        ('"뭐 없어?"', '카테고리 질문', '"어떤 종류의 옷을 찾으시나요?"'),
        ('"그냥 아무거나"', '추천 모드', '"요즘 인기 있는 상품 보여드릴까요?"'),
        ('"갬성템"', 'LLM 해석', '신조어 → 스타일 추론'),
        ('"위에 입는 거"', '카테고리 추론', '"상의를 찾으시는 거죠?"'),
        ('"예쁜 거"', '구체화 요청', '"어떤 스타일이 마음에 드세요? 캐주얼/포멀/..."'),
    ]
    add_table(doc, '', ambiguous_inputs, ['입력', '처리 방식', '응답 예시'])

    doc.add_heading('4.4 오류 입력 복구', level=2)

    error_recovery = [
        ('오타', '"겁정 바지"', '자동 교정 → "검정 바지로 검색할게요"'),
        ('약어', '"나키", "아디"', '정규화 → "나이키", "아디다스"'),
        ('외래어', '"black pants"', '매핑 → "검정 바지"'),
        ('혼용', '"블랙색"', '정규화 → "black"'),
        ('불완전 문장', '"그거 ㄱ"', '확인 질문 → "그 상품을 담아드릴까요?"'),
    ]
    add_table(doc, '', error_recovery, ['유형', '입력 예시', '처리'])

    doc.add_page_break()

    # ================================================================
    # 5. 응답 생성 가이드라인
    # ================================================================
    doc.add_heading('5. 응답 생성 가이드라인', level=1)

    doc.add_heading('5.1 응답 구조', level=2)

    response_structure = '''
응답 구조:
┌─────────────────────────────────────────────────────────┐
│ 1. 핵심 정보 (필수)                                      │
│    - 요청에 대한 직접적인 답변                            │
│    - 상품 리스트, 피팅 이미지, 주문 상태 등               │
│                                                         │
│ 2. 부가 정보 (선택)                                      │
│    - 추천 이유, 스타일 팁, 매칭 점수 등                   │
│    - 필요한 경우에만 포함                                │
│                                                         │
│ 3. 후속 액션 제안 (권장)                                 │
│    - "피팅해볼까요?", "장바구니에 담을까요?"              │
│    - 자연스러운 다음 단계 유도                           │
│                                                         │
└─────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, response_structure)

    doc.add_heading('5.2 톤앤매너', level=2)

    tone_guidelines = [
        ('친근함', '존댓말 + 부드러운 어미', '"~할까요?", "~드릴게요"'),
        ('간결함', '핵심 먼저, 부연 나중', '3줄 이내 요약 후 상세'),
        ('적극성', '대안 제시', '"없으면 비슷한 것도 찾아드릴까요?"'),
        ('공감', '실패 시 사과', '"아쉽지만...", "죄송해요..."'),
        ('이모지', '최소한만 사용', '카테고리 구분, 강조 시에만'),
    ]
    add_table(doc, '', tone_guidelines, ['원칙', '방법', '예시'])

    doc.add_heading('5.3 액션 버튼 제안', level=2)

    doc.add_paragraph(
        '응답 끝에 자연스러운 후속 액션을 제안하여 사용자가 쉽게 다음 단계로 진행할 수 있도록 합니다.'
    )

    action_suggestions = [
        ('검색 후', '"피팅해볼까요?" / "다른 조건으로 찾아볼까요?"'),
        ('피팅 후', '"마음에 드시면 담아드릴까요?" / "다른 상품도 입어볼까요?"'),
        ('장바구니 담기 후', '"주문하시겠어요?" / "더 쇼핑하시겠어요?"'),
        ('오류 발생 시', '"다시 시도해볼까요?" / "다른 조건으로 찾아볼까요?"'),
    ]
    add_table(doc, '', action_suggestions, ['상황', '제안 예시'])

    doc.add_heading('5.4 오류 메시지', level=2)

    error_messages = [
        ('API 오류', '"잠시 문제가 생겼어요. 다시 시도해볼게요."', '자동 재시도 → 실패 시 안내'),
        ('검색 결과 없음', '"조건에 맞는 상품을 찾지 못했어요. X 조건을 빼볼까요?"', '조건 완화 제안'),
        ('품절', '"아쉽지만 이 상품은 품절이에요. 비슷한 상품을 찾아드릴까요?"', '대안 제시'),
        ('피팅 실패', '"피팅 처리 중 문제가 생겼어요. 다시 시도해볼까요?"', '재시도 제안'),
        ('권한 없음', '"이 기능은 로그인 후 이용할 수 있어요."', '로그인 유도'),
    ]
    add_table(doc, '', error_messages, ['상황', '응답 예시', '전략'])

    doc.add_page_break()

    # ================================================================
    # 6. Edge Case 처리 우선순위
    # ================================================================
    doc.add_heading('6. Edge Case 처리 우선순위', level=1)

    doc.add_heading('6.1 P0: 필수 (서비스 중단 방지)', level=2)

    p0_cases = [
        ('빈 입력', '텍스트/이미지 없음', '"무엇을 도와드릴까요?"'),
        ('손상 이미지', '읽기 불가', '"이미지를 읽을 수 없어요. 다시 올려주세요."'),
        ('부적절 이미지', 'NSFW', '"이 이미지는 처리할 수 없어요."'),
        ('API 타임아웃', 'Vision/Claude', '재시도 3회 → 에러 안내'),
        ('품절 상품', '재고 없음', '대안 상품 추천'),
        ('동시 재고 충돌', '1개 남은 상품 동시 주문', '선착순 처리'),
    ]
    add_table(doc, '', p0_cases, ['케이스', '상황', '처리'])

    doc.add_heading('6.2 P1: 핵심 UX (자연스러운 대화)', level=2)

    p1_cases = [
        ('참조 해석', '"이거", "1번", "아까 그거"', '세션 컨텍스트 조회'),
        ('멀티턴 필터', '조건 누적/교체', '대화 상태 관리'),
        ('검색 수정', '"검정으로", "더 싸게"', '필터 추가/변경'),
        ('피팅 요청', '"입어볼래", "비교해줘"', '전제조건 확인 후 실행'),
        ('장바구니', '"담아줘", "빼줘"', 'CRUD 처리'),
        ('확인/선택', '"응", "아니", "1번"', '대화 흐름 분기'),
        ('사이즈 추천', '"사이즈 뭐가 좋아?"', '프로필 기반 추천'),
    ]
    add_table(doc, '', p1_cases, ['케이스', '예시', '처리'])

    doc.add_heading('6.3 P2: 품질 향상', level=2)

    p2_cases = [
        ('이미지 품질', '흐릿함, 어두움', '분석 시도 + 품질 안내'),
        ('오타 교정', '"겁정", "나키"', '자동 교정'),
        ('복잡한 참조', '"아까 그거에서 색만 바꾼"', '참조 체인 해석'),
        ('조건 완화', '결과 없을 때', '자동 필터 완화'),
        ('세션 복구', '장시간 대기 후', '상태 복원'),
        ('피드백 처리', '"별로야", "다른 거"', '대안 제시'),
    ]
    add_table(doc, '', p2_cases, ['케이스', '상황', '처리'])

    doc.add_heading('6.4 P3: 추가 개선', level=2)

    p3_cases = [
        ('일반 대화', '"안녕", "고마워"', '인사/감사 응대'),
        ('범위 외 질문', '"날씨 어때?"', '서비스 안내'),
        ('캐시 최적화', '중복 이미지', '캐시 활용'),
        ('개인화', '선호도 학습', '추천 개선'),
    ]
    add_table(doc, '', p3_cases, ['케이스', '상황', '처리'])

    doc.add_page_break()

    # ================================================================
    # 7. 구현 로드맵
    # ================================================================
    doc.add_heading('7. 구현 로드맵', level=1)

    roadmap = [
        ('1단계', 'Orchestrator 기본', 'Intent 분류, 서브 에이전트 라우팅, 세션 관리', '필수'),
        ('2단계', '참조 해석', '"이거", "1번" 등 참조 표현 처리', '필수'),
        ('3단계', '멀티턴 대화', '조건 누적, 교체, 대화 상태 관리', '필수'),
        ('4단계', '검색 에이전트', '이미지 분석, 텍스트 검색, 필터링', '필수'),
        ('5단계', '피팅 에이전트', '단일/배치/비교 피팅', '필수'),
        ('6단계', '커머스 에이전트', '사이즈 추천, 장바구니, 주문', '필수'),
        ('7단계', '복합 처리', '순차 체인, 조건부 실행', '권장'),
        ('8단계', '오류 복구', '오타 교정, 조건 완화', '권장'),
        ('9단계', '품질 개선', '응답 품질, 개인화', '선택'),
    ]
    add_table(doc, '', roadmap, ['단계', '범위', '내용', '우선순위'])

    # 의존성 다이어그램
    dependency_diagram = '''
구현 의존성:
┌─────────────────────────────────────────────────────────────────┐
│                                                                 │
│   [1단계: Orchestrator]                                         │
│         │                                                       │
│         ▼                                                       │
│   [2단계: 참조 해석] ─────────┐                                  │
│         │                    │                                  │
│         ▼                    ▼                                  │
│   [3단계: 멀티턴] ←───── [4단계: Search Agent]                   │
│         │                    │                                  │
│         │                    ▼                                  │
│         │            [5단계: Fitting Agent]                     │
│         │                    │                                  │
│         │                    ▼                                  │
│         └──────────→ [6단계: Commerce Agent]                    │
│                              │                                  │
│                              ▼                                  │
│                      [7단계: 복합 처리]                          │
│                              │                                  │
│                    ┌─────────┴─────────┐                        │
│                    ▼                   ▼                        │
│            [8단계: 오류 복구]  [9단계: 품질 개선]                 │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘
'''
    add_code_block(doc, dependency_diagram, '의존성 다이어그램:')

    doc.add_page_break()

    # ================================================================
    # 8. 부록: 응답 예시 모음
    # ================================================================
    doc.add_heading('부록: 응답 예시 모음', level=1)

    doc.add_heading('A. 검색 응답 예시', level=2)

    search_response_example = '''
👤 입력: [셔츠 이미지 첨부] "이거랑 어울리는 바지 찾아줘"

🤖 응답:
"흰색 스트라이프 셔츠와 잘 어울리는 바지를 찾았어요:

1. 네이비 치노팬츠 - ₩59,000
   클래식한 조합, 깔끔한 오피스룩

2. 베이지 슬랙스 - ₩45,000
   부드러운 느낌, 봄/가을 추천

3. 블랙 데님 - ₩69,000
   캐주얼 데일리룩

피팅해서 확인해볼까요?"
'''
    add_code_block(doc, search_response_example)

    doc.add_heading('B. 피팅 응답 예시', level=2)

    fitting_response_example = '''
👤 입력: "1번 입어볼래"

🤖 응답:
[피팅 이미지]

"네이비 치노팬츠 피팅 결과예요!
셔츠와 깔끔하게 매칭되네요.

컬러 매칭 점수: 92/100
추천 사이즈: 32인치 (평소 사이즈 기준)

장바구니에 담을까요?"
'''
    add_code_block(doc, fitting_response_example)

    doc.add_heading('C. 커머스 응답 예시', level=2)

    commerce_response_example = '''
👤 입력: "응 담아줘"

🤖 응답:
"장바구니에 담았어요! ✓

🛒 현재 장바구니:
┌─────────────────────────────────────┐
│ 네이비 치노팬츠 (32) × 1   ₩59,000 │
├─────────────────────────────────────┤
│ 합계                       ₩59,000 │
└─────────────────────────────────────┘

셔츠도 같이 찾아드릴까요?"
'''
    add_code_block(doc, commerce_response_example)

    doc.add_heading('D. 오류 응답 예시', level=2)

    error_response_example = '''
👤 입력: "빨간 구두 찾아줘"

🤖 응답 (결과 없음):
"아쉽지만 빨간 구두를 찾지 못했어요.

이렇게 해볼까요?
- 다른 색상으로 찾아보기 (검정/흰색/베이지)
- 빨간색 다른 신발 보기 (운동화/샌들)

어떤 것으로 찾아드릴까요?"
'''
    add_code_block(doc, error_response_example)

    # ================================================================
    # 저장
    # ================================================================
    output_path = '/Users/ijeong/Desktop/테커/AI_패션_어시스턴트_에이전트_개발_계획서.docx'
    doc.save(output_path)
    print(f"UX 중심 개발 계획서 저장: {output_path}")
    print(f"파일 크기: {os.path.getsize(output_path):,} bytes")


if __name__ == '__main__':
    create_ux_focused_plan()
